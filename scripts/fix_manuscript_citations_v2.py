"""
fix_manuscript_citations_v2.py
================================
Follow-up fixes to the manuscript after a user citation/number audit found:

1. The Moran's I subsection named the method but never cited it (no Moran
   1950). Fixed: adds "(Moran, 1950)" at first mention.
2. The DAG paragraph named "directed acyclic graph (DAG)" but never cited
   the standard epidemiology DAG-methodology reference. Fixed: adds
   "(Greenland, Pearl & Robins, 1999)".
3. The Hausman-test paragraph cited Hausman (1978) and Swamy & Arora
   (1972) but reference 14 (Wooldridge, 2010) -- the textbook this
   session's Swamy-Arora/Hausman implementation was actually built from
   (see scripts/robustness_hausman.py's docstring) -- was never cited
   in-text at all. Fixed: adds "; Wooldridge, 2010" alongside the Hausman
   citation.
4. Reference 13 (Cohen, 1988) is now ORPHANED (cited nowhere in text)
   because the E-value calculation was corrected (see
   scripts/robustness_dag_evalue.py) to no longer use a Cohen's-d
   intermediate step. Removed from the reference list.
5. The E-value paragraph itself is rewritten with the corrected number
   (3.9, not 2,093) and corrected methodology description.
6. References renumbered 9-15 to fit: 9 Cameron/Gelbach/Miller, 10
   Hausman, 11 Swamy & Arora, 12 VanderWeele & Ding, 13 Wooldridge
   (kept, now cited), 14 Greenland/Pearl/Robins (new), 15 Moran (new).

Run once against the current manuscript state (after update_manuscript.py
and fix_manuscript_formatting.py have already been applied).
"""

import docx
from docx.shared import Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH

PATH = "outputs/paper/PM25_Pediatric_Asthma_Philippines_REFORMATTED.docx"
FONT = "Times New Roman"

doc = docx.Document(PATH)


def find_para(text_exact=None, text_startswith=None, text_contains=None):
    for p in doc.paragraphs:
        t = p.text.strip()
        if text_exact is not None and t == text_exact:
            return p
        if text_startswith is not None and t.startswith(text_startswith):
            return p
        if text_contains is not None and text_contains in t:
            return p
    raise ValueError(f"paragraph not found: {text_exact or text_startswith or text_contains}")


def style_run(run, size, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return run


# ─────────────────────────────────────────────────────────────────────────
# 1. Moran's I citation
# ─────────────────────────────────────────────────────────────────────────
morans_para = find_para(text_contains="We tested this using Moran")
for r in morans_para.runs:
    if "We tested this using Moran" in r.text:
        r.text = r.text.replace(
            "We tested this using Moran’s I on the primary model’s residuals",
            "We tested this using Moran’s I (Moran, 1950) on the primary model’s residuals",
        )
print("Fix 1 applied: Moran (1950) citation added.")

# ─────────────────────────────────────────────────────────────────────────
# 2. DAG citation
# ─────────────────────────────────────────────────────────────────────────
dag_para = find_para(text_startswith="Figure 9 lays out the argument")
for r in dag_para.runs:
    if "directed acyclic graph (DAG)" in r.text:
        r.text = r.text.replace(
            "directed acyclic graph (DAG): region",
            "directed acyclic graph (DAG; Greenland, Pearl & Robins, 1999): region",
        )
print("Fix 2 applied: Greenland, Pearl & Robins (1999) citation added.")

# ─────────────────────────────────────────────────────────────────────────
# 3. Wooldridge citation alongside Hausman
# ─────────────────────────────────────────────────────────────────────────
hausman_para = find_para(text_contains="via a Hausman specification test")
for r in hausman_para.runs:
    if "via a Hausman specification test (Hausman, 1978)" in r.text:
        r.text = r.text.replace(
            "via a Hausman specification test (Hausman, 1978).",
            "via a Hausman specification test (Hausman, 1978; Wooldridge, 2010).",
        )
print("Fix 3 applied: Wooldridge (2010) now cited in-text.")

# ─────────────────────────────────────────────────────────────────────────
# 4. Rewrite the E-value paragraph with the corrected number/methodology
# ─────────────────────────────────────────────────────────────────────────
evalue_para = find_para(text_startswith="As a complementary, approximate quantification")
assert len(evalue_para.runs) >= 1
# Clear all runs' text, then set the first run to the full corrected text
# (keeps the paragraph's own formatting/position; simplest reliable way to
# replace a whole python-docx paragraph's content without re-inserting).
for i, r in enumerate(evalue_para.runs):
    r.text = "" if i > 0 else (
        "As a complementary, approximate quantification of how strong an unmeasured confounder "
        "would need to be, we calculated an E-value (VanderWeele & Ding, 2017) for the pooled "
        "cross-sectional correlation (r = +0.887). The E-value framework is defined for risk "
        "ratios, not Pearson correlations, so this used VanderWeele & Ding’s (2017) own stated "
        "approximation for a continuous, standardized effect size (their Table 2): "
        "RR ≈ exp(0.91 × d). For a bivariate relationship between two continuous, "
        "standardized variables, the standardized regression coefficient is the Pearson "
        "correlation itself (d = r = 0.887 here), so no further conversion (e.g. through an "
        "odds ratio) was needed. This gives RR ≈ 2.24 and an E-value ≈ 3.9: an "
        "unmeasured confounder would need to be associated with both PM2.5 and asthma "
        "prevalence by a risk ratio of at least roughly 3.9, above and beyond region and year, "
        "to fully explain away the pooled association on its own. This should be read as an "
        "approximation, not a precise bound — VanderWeele & Ding note this conversion carries "
        "modest error for typical effect sizes and more for very large ones, and r = 0.887 is a "
        "very large one — and it is computed on the pooled/confounded correlation, not on the "
        "FE-adjusted null result (which has no correlation coefficient to convert). An E-value "
        "of this size is a moderate, not extreme, bar. It is consistent with — not evidence "
        "for — this manuscript’s central argument: region-level confounding (urbanization, "
        "diagnostic capacity) of roughly this magnitude is independently identified and "
        "adjusted for via fixed effects, which is exactly why the pooled association collapses "
        "under the two-way FE model."
    )
    style_run(r, 11, bold=False, italic=False)
print("Fix 4 applied: E-value paragraph rewritten with corrected number (~3.9, not ~2,093).")

# ─────────────────────────────────────────────────────────────────────────
# 5. Reference list: remove Cohen (13), keep+renumber Wooldridge (was 14),
#    insert Greenland/Pearl/Robins and Moran, renumber 9-15
# ─────────────────────────────────────────────────────────────────────────
cohen_para = find_para(text_startswith="13. Cohen J. Statistical Power Analysis")
wooldridge_para = find_para(text_startswith="14. Wooldridge JM.")

# Renumber Wooldridge from 14 -> 13
for r in wooldridge_para.runs:
    if r.text.strip().startswith("14. Wooldridge"):
        r.text = r.text.replace("14. Wooldridge", "13. Wooldridge", 1)

# Delete Cohen paragraph (remove its XML element entirely)
cohen_para._element.getparent().remove(cohen_para._element)

new_refs = [
    "14. Greenland S, Pearl J, Robins JM. Causal diagrams for epidemiologic research. "
    "Epidemiology. 1999;10(1):37-48.",
    "15. Moran PAP. Notes on continuous stochastic phenomena. Biometrika. 1950;37(1-2):17-23.",
]
# insert AFTER wooldridge_para: python-docx has no insert_paragraph_after,
# so anchor on the paragraph that currently follows it (the AI Disclosure
# heading, same anchor update_manuscript.py used) and insert before that.
ai_disclosure_h1 = find_para(text_exact="Artificial Intelligence Disclosure")
for ref_text in new_refs:
    p = ai_disclosure_h1.insert_paragraph_before("")
    p.paragraph_format.space_after = Emu(88900)
    p.paragraph_format.left_indent = Emu(273685)
    p.paragraph_format.first_line_indent = Emu(-273685)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style_run(p.add_run(ref_text), 11, bold=False)

print("Fix 5 applied: reference list renumbered 9-15 (Cohen removed, "
      "Greenland/Pearl/Robins + Moran added, Wooldridge renumbered and now cited).")

doc.save(PATH)
print("\nSaved", PATH)
