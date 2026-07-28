"""
task4_mediator_manuscript_addition.py
=======================================
Writes the NDHS biomass-fuel-use mediator analysis (scripts/task3_mediator_
analysis.py, independently confirmed via real linearmodels.PanelOLS in
scripts/verify_with_linearmodels.py section 7) into the manuscript, as a
new subsection in the Discussion, right after the E-value discussion and
before "Why Asthma Prevalence Is the Wrong Outcome Variable" -- this is
the natural home because the E-value paragraph is exactly the qualitative/
theoretical confounding bound this new subsection complements with a
direct, data-based test of one candidate confounder.

Reported as EXPLORATORY, not as a confirmed finding, because the result is
not robust to dropping NCR alone (see task3_mediator_analysis.py output):
this is stated plainly in the new text itself, not left for a reader to
discover only by checking outputs/tables/mediator_biomass_fuel_results.csv.
Not added to the Abstract (unlike the Davao version-drift finding in
task1_davao_reframe.py) because an exploratory, NCR-fragile check does not
rise to the level of a named contribution.
"""

import docx
from docx.shared import Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH

PATH = "outputs/paper/PM25_Pediatric_Asthma_Philippines_REFORMATTED.docx"
FONT = "Times New Roman"

doc = docx.Document(PATH)


def find_para(text_exact=None, text_startswith=None, occurrence=1):
    count = 0
    for p in doc.paragraphs:
        t = p.text.strip()
        match = (text_exact is not None and t == text_exact) or \
                (text_startswith is not None and t.startswith(text_startswith))
        if match:
            count += 1
            if count == occurrence:
                return p
    raise ValueError(f"paragraph not found: {text_exact or text_startswith} (occurrence {occurrence})")


def style_run(run, size=11, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return run


def insert_h2(anchor, text):
    p = anchor.insert_paragraph_before("")
    p.paragraph_format.space_before = Pt(11)
    p.paragraph_format.space_after = Pt(6)
    style_run(p.add_run(text), 11, bold=True)
    return p


def insert_body(anchor, text):
    p = anchor.insert_paragraph_before("")
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style_run(p.add_run(text), 11, bold=False)
    return p


word_count_before = sum(len(p.text.split()) for p in doc.paragraphs)

# ---------------------------------------------------------------------
# 1. New H2 subsection + 3 paragraphs, inserted right before "Why Asthma
#    Prevalence Is the Wrong Outcome Variable" (i.e. right after the
#    E-value paragraph, which this subsection directly complements).
# ---------------------------------------------------------------------
anchor = find_para(text_exact="Why Asthma Prevalence Is the Wrong Outcome Variable")

insert_h2(anchor, "Exploratory Check: A Candidate Mediator for the Between-Region Confounding Structure")

insert_body(anchor,
    "As a direct empirical complement to the confounding argument and E-value bound above, indoor "
    "solid/biomass-fuel use for cooking — a plausible proxy for the region-level urbanization and "
    "socioeconomic differences already hypothesized to drive the pooled correlation — was tested as "
    "a candidate mediator. Region-level estimates were assembled from the three Philippine National "
    "Demographic and Health Survey (NDHS) rounds that overlap this study's window (2013, 2017, 2022; "
    "Philippine Statistics Authority & ICF, 2014, 2023; Wang et al., 2020), the only years with a "
    "fuel-use module. The 2022 estimate is reported directly at the region level; the 2017 estimate "
    "is aggregated from 81 province-level estimates (Wang et al., 2020) using the same "
    "province-to-region mapping used elsewhere in this study; the 2013 survey published no "
    "region-level breakdown at all, so its regional values were instead constructed as a composite "
    "of the national urban (38.7%) and rural (81.1%) rates, weighted by each region's own "
    "urban/rural household share from the same survey's sampling frame — an approximation, not a "
    "direct measurement (see Limitations). The three anchor years were linearly interpolated to fill "
    "the remaining study years, without extrapolating beyond 2013 or 2022."
)

insert_body(anchor,
    "Across the 17 regions, biomass-fuel-use correlates with PM2.5 at r = -0.73 (negative because "
    "urbanized, higher-outdoor-PM2.5 regions rely least on solid cooking fuel) — consistent with it "
    "proxying the regional structure discussed above. Adding it as a covariate to the primary "
    "two-way fixed-effects model moves the PM2.5 coefficient from β = -2.554 (p = 0.002) to "
    "β = -2.096 (p = 0.018), an 18% reduction; the biomass-fuel-use coefficient itself is not "
    "significant (β = 0.367, p = 0.168). This is consistent with some overlap between "
    "biomass-fuel-use and the PM2.5-asthma association, though the reduction falls well short of "
    "explaining the pooled-versus-within-region gap this study attributes to the confounding "
    "structure above."
)

insert_body(anchor,
    "This result should be read cautiously: it is not robust to the National Capital Region (NCR) "
    "alone. NCR already carries disproportionate weight in the primary result (dropping it alone "
    "widens the primary p-value from 0.002 to 0.045; see Leave-One-Region-Out Jackknife above), and "
    "it is also the region where the 2013 composite estimate is least reliable — NCR's fully-urban "
    "household frame pins its 2013 composite to the national urban rate (38.7%) alone, far above its "
    "real 2017 (4.0%) and 2022 (1.2%) survey values. Dropping NCR alone reduces the between-region "
    "correlation to r = -0.42 and makes the covariate-adjusted PM2.5 coefficient non-significant "
    "(p = 0.15). Given both the dataset's general NCR-sensitivity and this specific weakness in the "
    "2013 estimate, this check is reported as exploratory and suggestive of a possible confound, not "
    "as a confirmed mediation finding."
)
print("Fix 1 applied: new 'Exploratory Check' subsection (3 paragraphs) inserted after the E-value discussion.")

# ---------------------------------------------------------------------
# 2. Limitations: one short cross-referencing sentence -- doesn't
#    duplicate the caveat above, just points to it so a reader scanning
#    Limitations doesn't miss that this WAS tested (partially).
# ---------------------------------------------------------------------
limitations = find_para(text_startswith="Several limitations must be acknowledged.")
old_seg = ("No confounders beyond region and year effects were included; residual confounding from "
           "healthcare access, smoking, and socioeconomic factors may remain.")
new_seg = (old_seg + " One candidate confounder — indoor biomass-fuel use, a proxy for regional "
           "urbanization — was tested directly (see Discussion above) and showed a modest association "
           "that is not robust to a single region (NCR), so this limitation is only partially addressed.")
hit = False
for r in limitations.runs:
    if old_seg in r.text:
        r.text = r.text.replace(old_seg, new_seg)
        hit = True
if not hit:
    raise ValueError("Limitations sentence not found for the cross-reference addition")
print("Fix 2 applied: Limitations cross-references the new exploratory mediator check.")

# ---------------------------------------------------------------------
# 3. Reference list -- add 3 new entries (25-27) for the NDHS/DHS
#    sources, matching the existing hanging-indent paragraph formatting
#    used for reference 24 (Holm, 1979) in task2_multiplicity.py.
# ---------------------------------------------------------------------
ai_disclosure_h1 = find_para(text_exact="Artificial Intelligence Disclosure")

new_refs = [
    "25. Philippine Statistics Authority (PSA) and ICF. Philippines National Demographic and "
    "Health Survey 2013. Manila, Philippines, and Rockville, Maryland, USA: PSA and ICF; 2014.",
    "26. Wang W, Assaf S, Mayala B, Moonzwe Davis L. Household Air Pollution: National and "
    "Subnational Estimates in Bangladesh, India, Indonesia, Nepal, and the Philippines. DHS "
    "Working Paper No. 164. Rockville, Maryland, USA: ICF; 2020.",
    "27. Philippine Statistics Authority (PSA) and ICF. 2022 Philippines National Demographic "
    "and Health Survey (NDHS) Final Report. Quezon City, Philippines, and Rockville, Maryland, "
    "USA: PSA and ICF; 2023.",
]
for ref_text in new_refs:
    new_ref = ai_disclosure_h1.insert_paragraph_before("")
    new_ref.paragraph_format.left_indent = Emu(273685)
    new_ref.paragraph_format.first_line_indent = Emu(-273685)
    new_ref.paragraph_format.space_after = Emu(88900)
    style_run(new_ref.add_run(ref_text), 11, bold=False)
print("Fix 3 applied: references 25-27 (NDHS 2013, WP164, NDHS 2022) added.")

doc.save(PATH)

word_count_after = sum(len(p.text.split()) for p in doc.paragraphs)
print(f"\nWord count before: {word_count_before}")
print(f"Word count after:  {word_count_after}")
print(f"Net word change:   {word_count_after - word_count_before:+d}")
print(f"\nSaved {PATH}")
