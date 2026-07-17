"""
rebuild_reference_list.py
===========================
Removes the current references 9-15 and replaces them with the full,
final 9-23 list: the same 7 references (now with user-supplied DOIs added)
plus the 8 genuinely new ones from this round, ordered to match the order
their subjects appear in the manuscript (WCR bootstrap -> Hausman/RE ->
MDE -> negative controls -> population weighting -> E-value -> DAG ->
Moran's I -> ecological fallacy).

DOI formatting matches the majority existing convention in this reference
list: lowercase "doi:10.xxxx" with no space (refs 1, 2, 8 use this; only
ref 7 uses "DOI: 10.xxxx" uppercase-with-space, so lowercase-no-space is
the majority pattern and is used for every DOI added/added-to below).

Run AFTER add_15_references.py (which handles the in-text citations).
"""

import docx
from docx.shared import Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH

PATH = "outputs/paper/PM25_Pediatric_Asthma_Philippines_REFORMATTED.docx"
FONT = "Times New Roman"

doc = docx.Document(PATH)


def find_para(text_exact=None, text_startswith=None):
    for p in doc.paragraphs:
        t = p.text.strip()
        if text_exact is not None and t == text_exact:
            return p
        if text_startswith is not None and t.startswith(text_startswith):
            return p
    raise ValueError(f"paragraph not found: {text_exact or text_startswith}")


def style_run(run, size=11, bold=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    return run


# ─────────────────────────────────────────────────────────────────────────
# 1. DELETE current refs 9-15 (7 paragraphs)
# ─────────────────────────────────────────────────────────────────────────
old_ref_starts = [f"{n}. " for n in range(9, 16)]
to_delete = []
for p in doc.paragraphs:
    t = p.text.strip()
    if any(t.startswith(s) for s in old_ref_starts):
        to_delete.append(p)
assert len(to_delete) == 7, f"expected 7 old references (9-15), found {len(to_delete)}"
anchor_after = to_delete[0]  # insert the new block right where ref 9 used to start
for p in to_delete:
    p._element.getparent().remove(p._element)

ai_disclosure_h1 = find_para(text_exact="Artificial Intelligence Disclosure")

# ─────────────────────────────────────────────────────────────────────────
# 2. FINAL 9-23 REFERENCE LIST (order matches subject order in the manuscript)
# ─────────────────────────────────────────────────────────────────────────
new_refs = [
    "9. Cameron AC, Gelbach JB, Miller DL. Bootstrap-based improvements for inference "
    "with clustered errors. The Review of Economics and Statistics. 2008;90(3):414-427. "
    "doi:10.1162/rest.90.3.414",

    "10. MacKinnon JG, Nielsen MØ, Webb MD. Cluster-robust inference: a guide to "
    "empirical practice. Journal of Econometrics. 2023;232(2):272-299.",

    "11. Hausman JA. Specification tests in econometrics. Econometrica. 1978;46(6):"
    "1251-1271. doi:10.2307/1913827",

    "12. Swamy PAVB, Arora SS. The exact finite sample properties of the estimators of "
    "coefficients in the error components regression models. Econometrica. 1972;40(2):"
    "261-275.",

    "13. Wooldridge JM. Econometric Analysis of Cross Section and Panel Data. 2nd ed. "
    "Cambridge, MA: MIT Press; 2010.",

    "14. Bloom HS. Minimum detectable effects: a simple way to report the statistical "
    "power of experimental designs. Evaluation Review. 1995;19(5):547-556. "
    "doi:10.1177/0193841X9501900504",

    "15. Lipsitch M, Tchetgen Tchetgen E, Cohen T. Negative controls: a tool for "
    "detecting confounding and bias in observational studies. Epidemiology. 2010;"
    "21(3):383-388. doi:10.1097/EDE.0b013e3181d61eeb",

    "16. Ivy D, Mulholland JA, Russell AG. Development of ambient air quality "
    "population-weighted metrics for use in time-series health studies. Journal of "
    "the Air & Waste Management Association. 2008;58(5):711-720. "
    "doi:10.3155/1047-3289.58.5.711",

    "17. VanderWeele TJ, Ding P. Sensitivity analysis in observational research: "
    "introducing the E-value. Annals of Internal Medicine. 2017;167(4):268-274. "
    "doi:10.7326/M16-2607",

    "18. Haneuse S, VanderWeele TJ, Arterburn D. Using the E-value to assess the "
    "potential effect of unmeasured confounding in observational studies. JAMA. "
    "2019;321(6):602-603. doi:10.1001/jama.2018.21554",

    "19. Greenland S, Pearl J, Robins JM. Causal diagrams for epidemiologic research. "
    "Epidemiology. 1999;10(1):37-48. doi:10.1097/00001648-199901000-00008",

    "20. Textor J, van der Zander B, Gilthorpe MS, Liśkiewicz M, Ellison GTH. Robust "
    "causal inference using directed acyclic graphs: the R package 'dagitty'. "
    "International Journal of Epidemiology. 2016;45(6):1887-1894. doi:10.1093/ije/dyw341",

    "21. Moran PAP. Notes on continuous stochastic phenomena. Biometrika. 1950;"
    "37(1-2):17-23. doi:10.1093/biomet/37.1-2.17",

    "22. Anselin L. Local indicators of spatial association—LISA. Geographical "
    "Analysis. 1995;27(2):93-115. doi:10.1111/j.1538-4632.1995.tb00338.x",

    "23. Robinson WS. Ecological correlations and the behavior of individuals. "
    "American Sociological Review. 1950;15(3):351-357. doi:10.2307/2087176",
]

for ref_text in new_refs:
    p = ai_disclosure_h1.insert_paragraph_before("")
    p.paragraph_format.space_after = Emu(88900)
    p.paragraph_format.left_indent = Emu(273685)
    p.paragraph_format.first_line_indent = Emu(-273685)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style_run(p.add_run(ref_text))

print(f"Inserted {len(new_refs)} references (9-23), replacing the old 9-15 block.")

doc.save(PATH)
print("Saved", PATH)
