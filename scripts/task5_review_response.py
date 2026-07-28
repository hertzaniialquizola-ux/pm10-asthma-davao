"""
task5_review_response.py
==========================
Implements the user-approved response to an external methods-focused peer
review: (1) a new Methods paragraph naming the identification strategy and
scoping the region-trends check into the uncorrected-checks bucket, (2) a
new, prominent "Region-Specific Linear Time Trends" Results subsection
(the identification-relevant check), (3) collapsing 4 SE-reliability
checks (permutation, WCR, jackknife, Moran's I) into one compact section +
display table (keeping the existing jackknife figure), (4) a new
variance-decomposition-across-6-outcomes display table + a precision-
corrected rewrite of the "stock measure" paragraph, (5) an updated
Synthesis of Sensitivity Analyses paragraph, (6) an updated Abstract/
Conclusion thesis, (7) a new reference (Wolfers, 2006 -- verified via web
search, not memory, per this project's standing citation rule).

NOTE ON TABLE NUMBERING: the two new display tables (SE-reliability
summary, 6-outcome variance decomposition) are given bold descriptive
titles WITHOUT a sequential "Table N." caption, specifically to avoid a
high-risk mass renumbering of the 10 existing "Table 4/5/6" references
scattered across the manuscript (a blind find-replace across
run-fragmented text risked silently corrupting references in a
submission-bound document). Flagged for the user's awareness, not decided
silently -- easy to convert to numbered tables later if wanted.
"""

import docx
from docx.shared import Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy

PATH = "outputs/paper/PM25_Pediatric_Asthma_Philippines_REFORMATTED.docx"
FONT = "Times New Roman"

doc = docx.Document(PATH)


def find_para(text_exact=None, text_startswith=None, occurrence=1):
    count = 0
    for p in doc.paragraphs:
        t = p.text.strip()
        match = (text_exact is not None and t == text_exact) or \
                (text_startswith is not None and t.startswith(text_startswith))
        if match:
            count += 1
            if count == occurrence:
                return p
    raise ValueError(f"paragraph not found: {text_exact or text_startswith} (occurrence {occurrence})")


def style_run(run, size=11, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return run


def insert_h2(anchor, text):
    p = anchor.insert_paragraph_before("")
    p.paragraph_format.space_before = Pt(11)
    p.paragraph_format.space_after = Pt(6)
    style_run(p.add_run(text), 11, bold=True)
    return p


def insert_body(anchor, text):
    p = anchor.insert_paragraph_before("")
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style_run(p.add_run(text), 11, bold=False)
    return p


def insert_bold_lead(anchor, text):
    """A bolded, italicized mini-caption for an unnumbered display table."""
    p = anchor.insert_paragraph_before("")
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    style_run(p.add_run(text), 10.5, bold=True, italic=True)
    return p


def replace_paragraph_text(para, new_text):
    runs = para.runs
    if not runs:
        r = para.add_run(new_text)
        style_run(r)
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""


def delete_paragraph(para):
    p = para._element
    p.getparent().remove(p)


def clone_table_before(anchor_para, source_table, data_rows):
    """
    Deep-copy source_table's XML (preserving all borders/shading/font
    formatting), resize to len(data_rows)+1 rows (adding via add_row() if
    needed, removing extra rows if the source had more), fill in new text,
    and move the cloned table's XML element to sit immediately before
    anchor_para. `data_rows` is a list of lists of strings; row 0 of the
    ORIGINAL table is reused as the header row template and overwritten
    with data_rows[0] (the new header).
    """
    new_tbl_element = copy.deepcopy(source_table._tbl)
    new_table = docx.table.Table(new_tbl_element, source_table._parent)

    n_needed = len(data_rows)
    n_have = len(new_table.rows)
    while len(new_table.rows) < n_needed:
        new_table.add_row()
    while len(new_table.rows) > n_needed:
        last_tr = new_table.rows[-1]._tr
        last_tr.getparent().remove(last_tr)

    for r_idx, row_vals in enumerate(data_rows):
        for c_idx, val in enumerate(row_vals):
            cell = new_table.rows[r_idx].cells[c_idx]
            # Clear existing paragraphs' runs, keep first paragraph, set text.
            for p in cell.paragraphs[1:]:
                p._element.getparent().remove(p._element)
            p0 = cell.paragraphs[0]
            for r in p0.runs[1:]:
                r.text = ""
            if p0.runs:
                p0.runs[0].text = val
                style_run(p0.runs[0], 10.5, bold=(r_idx == 0))
            else:
                style_run(p0.add_run(val), 10.5, bold=(r_idx == 0))

    # Move the cloned table to sit right before anchor_para in the XML tree.
    anchor_para._element.addprevious(new_tbl_element)
    return new_table


word_count_before = sum(len(p.text.split()) for p in doc.paragraphs)

# =========================================================================
# 1. METHODS -- "Identification Strategy and Its Limits"
# =========================================================================
results_h1 = find_para(text_exact="Results")
insert_h2(results_h1, "Identification Strategy and Its Limits")
insert_body(results_h1,
    "This design identifies β from within-region, year-to-year covariation between PM2.5 and "
    "asthma prevalence, after removing all time-invariant regional differences (region fixed "
    "effects) and all shared national-level shocks (year fixed effects). The core identifying "
    "assumption is that, conditional on these two sets of effects, no remaining time-varying "
    "confounder is correlated with both PM2.5 and asthma prevalence within a region — most "
    "plausibly threatened by differential regional trajectories in healthcare access or "
    "diagnostic capacity (for example, the staggered rollout of the Universal Health Care Act "
    "and PhilHealth expansion across regions), rather than by differences in regional levels, "
    "which region fixed effects already remove. Of the robustness checks reported below, four "
    "(the permutation tests, wild cluster bootstrap, leave-one-region-out jackknife, and "
    "Moran’s I test) evaluate whether the standard error on β is trustworthy given the small "
    "number of clusters (17 regions); they do not test whether β itself could be confounded by "
    "a differential regional trajectory. The region-specific linear time trends check does test "
    "that directly, by allowing each region its own secular trend in addition to its own level. "
    "Consistent with the same logic already applied to the leave-one-region-out jackknife and "
    "the COVID-19 exclusion check — testing the stability of one pre-specified estimate rather "
    "than searching across candidate specifications — the region-specific trends check is not "
    "included in either Holm-Bonferroni-corrected family reported in the Results (the "
    "five-outcome comparison or the lag-structure sweep) and is reported at its raw "
    "significance level, consistent with those two checks."
)
print("Fix 1 applied: Methods 'Identification Strategy and Its Limits' paragraph inserted.")

# =========================================================================
# 2. RESULTS -- new "Region-Specific Linear Time Trends" subsection,
#    inserted as the FIRST robustness check (before "First Differences"),
#    giving it the prominence the identification-relevant check deserves.
# =========================================================================
first_diff_h2 = find_para(text_exact="Robustness Check: First Differences")
insert_h2(first_diff_h2, "Robustness Check: Region-Specific Linear Time Trends")
insert_body(first_diff_h2,
    "The checks that follow this one test whether the standard error on the primary "
    "coefficient is trustworthy given a modest number of clusters (see Identification Strategy "
    "and Its Limits, Methods); this check instead tests the identifying assumption itself. We "
    "re-estimated the primary model with one additional linear time trend per region (17 extra "
    "parameters: regionᵢ × t) alongside the existing region and year fixed effects — a standard "
    "approach in applied panel econometrics for testing whether a result reflects genuine "
    "within-region co-movement or a differential regional trajectory that happens to correlate "
    "with each region’s own PM2.5 trend (e.g. Wolfers, 2006)."
)
insert_body(first_diff_h2,
    "Adding region-specific trends shrank the coefficient from β = −2.554 (p = 0.0024) to "
    "β = −0.937 (SE = 0.446, p = 0.038, n = 170) — a 63% reduction in magnitude that narrowly "
    "survives conventional significance. This is a demanding specification (18 parameters "
    "estimated from 170 observations; within-R² rises mechanically from 0.080 to 0.791, since "
    "the region-specific trends absorb most of the panel’s variance, leaving relatively little "
    "residual variation to identify β from), so this result is better read as a lower bound on "
    "how much of the primary coefficient could plausibly reflect differential regional "
    "trajectories rather than a precise re-estimate in its own right. Taken at face value, it "
    "shows the primary result is not fully robust to letting each region follow its own trend "
    "— a materially more fragile picture than the standard-error-focused checks below suggest "
    "on their own, and one that reinforces, rather than undermines, this manuscript’s existing "
    "caution against over-reading the primary coefficient’s sign or magnitude as a confirmed "
    "effect."
)
print("Fix 2 applied: new 'Region-Specific Linear Time Trends' subsection inserted (2 paragraphs).")

# =========================================================================
# 3. RESULTS -- collapse Permutation / WCR / Jackknife-prose / Moran's I
#    into one unified section + compact display table. Figure 7
#    (jackknife plot) and its caption are KEPT, just re-homed under the
#    new heading.
# =========================================================================
perm_h2 = find_para(text_exact="Robustness Check: Permutation Test")
insert_h2(perm_h2, "SE-Reliability Checks: Small-Cluster Inference and Spatial Autocorrelation")
insert_body(perm_h2,
    "The primary model clusters standard errors by region, but with only 17 clusters that "
    "correction can itself be unreliable (Cameron & Miller, 2015). Four checks address this "
    "directly — none change the point estimate itself, only the assessment of its precision "
    "— and are summarized together below rather than as four separate specifications."
)
insert_bold_lead(perm_h2, "SE-reliability checks, summarized (primary estimate: β = −2.554, n = 170 throughout):")

se_table_rows = [
    ["Check", "What it tests", "Key result", "Verdict"],
    ["Permutation tests\n(label & block, 5,000 iter. each)",
     "Randomization-inference p-value under H0: β = 0",
     "Label p = 0.0048; block p = 0.0008",
     "Confirms significance"],
    ["Wild cluster bootstrap\n(WCR, 5,000 reps)",
     "Small-G (17-cluster) sign-flip bootstrap p-value",
     "p = 0.036 (t = −3.10)",
     "Confirms significance, narrower margin"],
    ["Leave-one-region-out jackknife\n(17 refits, Figure 7)",
     "Stability of β across single-region exclusions",
     "Range β = −1.416 to −2.812, all p < 0.05; weakest excluding NCR (p = 0.045)",
     "Confirms significance; NCR-sensitive"],
    ["Moran’s I\n(k = 4 nearest-neighbor weights)",
     "Spatial autocorrelation of residuals (SE validity)",
     "I = −0.0017 (expected −0.063 under H0), p = 0.994",
     "No evidence of spatial autocorrelation"],
]
# Clone Table 5's structure (4 columns: "Year"/"Ground PM2.5"/"Satellite PM2.5 V6GL03"/"Difference"),
# the only existing table with exactly 4 columns.
source_4col = None
for t in doc.tables:
    header = [c.text.strip() for c in t.rows[0].cells]
    if len(header) == 4 and header[0] == "Year" and "Ground PM2.5" in header[1] and "Difference" in header[3]:
        source_4col = t
        break
if source_4col is None:
    raise RuntimeError("Could not find a 4-column source table to clone for the SE-reliability table.")
clone_table_before(source_4col, source_4col, se_table_rows)
print("Fix 3a applied: SE-reliability display table inserted (cloned formatting from the Davao 4-col table).")

# Now delete the OLD 4 subsections' headings + prose (figure/caption for
# jackknife is preserved -- only its H2 + explanatory paragraph go).
for heading_text in [
    "Robustness Check: Permutation Test",
    "Robustness Check: Wild Cluster Bootstrap",
    "Robustness Check: Leave-One-Region-Out Jackknife",
    "Robustness Check: Spatial Autocorrelation of Residuals (Moran’s I)",
]:
    h2 = find_para(text_exact=heading_text)
    # the very next sibling element is always that check's explanatory prose
    next_elem = h2._element.getnext()
    prose_text = docx.text.paragraph.Paragraph(next_elem, h2._parent).text[:60]
    next_elem.getparent().remove(next_elem)
    delete_paragraph(h2)
    print(f"    removed heading {heading_text!r} + prose starting {prose_text!r}")
print("Fix 3b applied: old Permutation/WCR/Jackknife-prose/Moran's I headings and prose removed "
      "(Figure 7 and its caption preserved).")

# =========================================================================
# 4. DISCUSSION -- "Why Asthma Prevalence Is the Wrong Outcome Variable":
#    precision-corrected rewrite + new 6-outcome variance display table.
# =========================================================================
p_stock = find_para(text_startswith="This reflects a fundamental feature of asthma prevalence")
replace_paragraph_text(p_stock,
    "Whether this near-total between-region variance share reflects the epidemiology of "
    "asthma as a slowly accumulating chronic condition, or the smoothing inherent to GBD’s "
    "subnational estimation procedure more generally, is directly checkable using the "
    "placebo/negative-control outcome (low back pain) already used elsewhere in this study "
    "(Robustness Check: Placebo Outcome). Its own within-region variance share is 38.0% — 26 "
    "times higher than asthma’s 1.5% — and its between-region share (62.0%) is far below "
    "asthma’s (98.5%). Extending the comparison to the four other GBD-modeled outcomes "
    "examined in this study (Table 4) shows a clear gradient tracking disease acuity rather "
    "than a uniform floor: chronic, slowly-progressing conditions (asthma, 1.5% within-region; "
    "COPD, 2.0%) sit at one end, and acute, fast-changing conditions (lower respiratory "
    "infection incidence, 82.8%; respiratory mortality, 88.6%) sit at the other, with the "
    "placebo in between. If GBD’s subnational modeling suppressed within-region variance "
    "uniformly regardless of the specific cause being modeled, the placebo outcome should show "
    "a variance split similar to asthma’s; it does not. This argues against uniform smoothing "
    "as the explanation and supports the stock-measure interpretation above, though it does "
    "not fully rule out the possibility that GBD’s own spatiotemporal modeling incorporates "
    "disease-course priors that themselves encode this same acuity gradient. Under either "
    "interpretation, outcome variables with higher-frequency signal — lower respiratory "
    "infection incidence, emergency-department visits, hospital admission rates — would "
    "better suit a panel of this length."
)
print("Fix 4a applied: 'stock measure' paragraph rewritten with the acuity-gradient precision fix.")

testing_additional_h2 = find_para(text_exact="Testing Additional Respiratory Outcomes")
insert_bold_lead(testing_additional_h2,
    "Within- and between-region variance share across six GBD-modeled outcomes, Philippines 2013–2022:")
variance_table_rows = [
    ["Outcome", "Within-region %", "Between-region %"],
    ["Asthma (primary outcome)", "1.5%", "98.5%"],
    ["COPD prevalence", "2.0%", "98.0%"],
    ["Lung cancer incidence", "6.5%", "93.5%"],
    ["Low back pain (placebo/negative control)", "38.0%", "62.0%"],
    ["Lower respiratory infection incidence", "82.8%", "17.2%"],
    ["Respiratory disease mortality", "88.6%", "11.4%"],
]
source_3col = None
for t in doc.tables:
    header = [c.text.strip() for c in t.rows[0].cells]
    if len(header) == 3 and header[0] == "Variable" and "Within-Region" in header[1]:
        source_3col = t
        break
if source_3col is None:
    raise RuntimeError("Could not find the 3-column variance-decomposition table to clone.")
clone_table_before(source_3col, source_3col, variance_table_rows)
print("Fix 4b applied: 6-outcome variance-decomposition display table inserted "
      "(cloned formatting from the existing Table 2).")

# =========================================================================
# 5. Synthesis of Sensitivity Analyses -- fold in the region-trends result.
# =========================================================================
synthesis = find_para(text_startswith="Across all robustness checks conducted for the primary coefficient")
old_open = ("Across all robustness checks conducted for the primary coefficient — clustered standard "
            "errors, label- and block-permutation tests, a wild cluster bootstrap, a leave-one-region-out "
            "jackknife, and exclusion of the COVID-19 pandemic years — the two-way fixed-effects estimate "
            "was directionally consistent (always negative) and statistically significant in every check "
            "except one (excluding 2020–2021, where it narrowly missed conventional significance at "
            "p = 0.088). The estimate was measurably, though not fully, sensitive to two specific features "
            "of the data: the National Capital Region and the two pandemic years.")
new_open = ("Across all robustness checks conducted for the primary coefficient — clustered standard "
            "errors, label- and block-permutation tests, a wild cluster bootstrap, a leave-one-region-out "
            "jackknife, region-specific linear time trends, and exclusion of the COVID-19 pandemic years — "
            "the two-way fixed-effects estimate was directionally consistent (always negative) and "
            "statistically significant in every check except one (excluding 2020–2021, where it narrowly "
            "missed conventional significance at p = 0.088). The estimate was measurably, and in one case "
            "substantially, sensitive to three specific features of the data: the National Capital Region, "
            "the two pandemic years, and — most consequentially — the possibility of differential regional "
            "trajectories rather than differential regional levels. Allowing each region its own linear "
            "time trend, in addition to its own fixed level, shrank the coefficient by 63% "
            "(β = −2.554 to β = −0.937) and left it only narrowly significant (p = 0.038); of every check "
            "in this manuscript, this is the one that most changes how confidently the primary result "
            "should be read, and it is treated here as real evidence of fragility rather than explained "
            "away.")
old_placebo_close = ("A placebo/negative-control test using an outcome with no plausible PM2.5 pathway "
                      "(low back pain prevalence) found no significant association (p = 0.664), supporting "
                      "the primary null as a genuine finding rather than an artifact of the GBD estimation "
                      "pipeline.")
new_placebo_close = ("A placebo/negative-control test using an outcome with no plausible PM2.5 pathway "
                      "(low back pain prevalence) found no significant association (p = 0.664), and its "
                      "own within-region variance share (38.0%, far above asthma’s 1.5%) argues against "
                      "uniform GBD-pipeline smoothing as the explanation for asthma’s near-total "
                      "between-region variance (Discussion above).")
old_tail = ("Taken together, these checks support treating the primary null/negative finding as a genuine "
            "feature of this dataset rather than an artifact of any single modeling choice, while also "
            "making explicit exactly where that finding is least secure (NCR, the pandemic years, and the "
            "overall statistical power of the design, quantified above as an MDE of ≈ 2.33).")
new_tail = ("Taken together, these checks support treating the primary null/negative finding as a genuine, "
            "but only narrowly robust, feature of this dataset rather than an artifact of any single "
            "modeling choice, while making explicit exactly where that finding is least secure: the "
            "National Capital Region, the pandemic years, the overall statistical power of the design "
            "(MDE ≈ 2.33), and, most of all, the possibility that differential regional trajectories "
            "rather than a shared within-region mechanism account for a substantial share of the primary "
            "coefficient’s magnitude.")

full_text = "".join(r.text for r in synthesis.runs)
if old_open not in full_text or old_placebo_close not in full_text or old_tail not in full_text:
    raise ValueError("Synthesis paragraph anchors not found exactly -- check for smart-quote mismatches.")
new_full_text = full_text.replace(old_open, new_open).replace(old_placebo_close, new_placebo_close).replace(old_tail, new_tail)
synthesis.runs[0].text = new_full_text
for r in synthesis.runs[1:]:
    r.text = ""
print("Fix 5 applied: Synthesis of Sensitivity Analyses updated with the region-trends result.")

doc.save(PATH)
print("\n[Checkpoint save after Fixes 1-5 -- continuing with Abstract/Conclusion/References in part 2]")

word_count_mid = sum(len(p.text.split()) for p in doc.paragraphs)
print(f"Word count before: {word_count_before}")
print(f"Word count after Fixes 1-5: {word_count_mid}")
