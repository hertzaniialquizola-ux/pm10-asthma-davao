"""
task5b_abstract_refs.py
=========================
Part 2 of the review-response edits: Abstract thesis update (the user's
tightened V2 wording, incorporating the region-specific-trends result),
Conclusion paragraph alignment, and a new reference (Wolfers, 2006 --
verified via web search before citing, per this project's standing rule
after a prior fabricated-citation incident).
"""

import docx
from docx.shared import Pt, Emu

PATH = "outputs/paper/PM25_Pediatric_Asthma_Philippines_REFORMATTED.docx"
FONT = "Times New Roman"

doc = docx.Document(PATH)


def find_para(text_exact=None, text_startswith=None, occurrence=1):
    count = 0
    for p in doc.paragraphs:
        t = p.text.strip()
        match = (text_exact is not None and t == text_exact) or \
                (text_startswith is not None and t.startswith(text_startswith))
        if match:
            count += 1
            if count == occurrence:
                return p
    raise ValueError(f"paragraph not found: {text_exact or text_startswith} (occurrence {occurrence})")


def style_run(run, size=11, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return run


word_count_before = sum(len(p.text.split()) for p in doc.paragraphs)

# =========================================================================
# 1. ABSTRACT -- update the Results sensitivity-analyses sentence to
#    mention the region-trends shrinkage, and replace the Conclusions
#    thesis sentence with the user's tightened V2 wording.
# =========================================================================
abstract = find_para(text_startswith="Background. Fine particulate matter")
full_text = "".join(r.text for r in abstract.runs)

old_sensitivity = ("Sensitivity analyses were mostly, but not uniformly, supportive: a wild "
    "cluster bootstrap (p = 0.036) and a leave-one-region-out jackknife (β ranged −1.416 to "
    "−2.812 across 17 refits, all p < 0.05) preserved the coefficient’s sign and significance, "
    "while excluding the COVID-19 pandemic years (2020–2021) attenuated it to non-significance "
    "(β = −2.113, p = 0.088) and a 3-year cumulative PM2.5 exposure specification instead "
    "produced a substantially larger, highly significant negative coefficient (β = −6.040, "
    "p < 0.0001).")
new_sensitivity = ("Sensitivity analyses were mostly, but not uniformly, supportive: a wild "
    "cluster bootstrap (p = 0.036) and a leave-one-region-out jackknife (β ranged −1.416 to "
    "−2.812 across 17 refits, all p < 0.05) preserved the coefficient’s sign and significance, "
    "while excluding the COVID-19 pandemic years (2020–2021) attenuated it to non-significance "
    "(β = −2.113, p = 0.088) and a 3-year cumulative PM2.5 exposure specification instead "
    "produced a substantially larger, highly significant negative coefficient (β = −6.040, "
    "p < 0.0001). Most consequentially, a specification allowing each region its own linear "
    "time trend — testing whether the result reflects differential regional trajectories "
    "rather than genuine within-region co-movement — shrank the coefficient by 63% "
    "(β = −0.937, p = 0.038).")

old_thesis = ("These sensitivity analyses show the fixed-effects estimate is reasonably, though "
    "not perfectly, robust: it is sensitive to the National Capital Region and to the "
    "COVID-19 pandemic years specifically, and strengthens rather than weakens under a more "
    "biologically plausible cumulative-exposure specification.")
new_thesis = ("Across nine robustness checks — including a specification allowing each region "
    "its own secular trend — the coefficient never reversed sign but shrank by as much as 63% "
    "and repeatedly sat at the edge of conventional significance, a pattern more consistent "
    "with a near-null relationship obscured by low within-region signal than with a robust "
    "negative or positive effect.")

if old_sensitivity not in full_text:
    raise ValueError("Abstract sensitivity-analyses sentence anchor not found.")
if old_thesis not in full_text:
    raise ValueError("Abstract thesis sentence anchor not found.")

new_full_text = full_text.replace(old_sensitivity, new_sensitivity).replace(old_thesis, new_thesis)
abstract.runs[0].text = new_full_text
for r in abstract.runs[1:]:
    r.text = ""
print("Fix 1 applied: Abstract sensitivity-analyses sentence + thesis sentence updated.")

# =========================================================================
# 2. CONCLUSION -- align the opening paragraph with the new thesis.
# =========================================================================
conclusion_open = find_para(text_startswith="This subnational ecological panel analysis")
old_c = ("After adjusting for region and year fixed effects, no robust association was "
    "identified, consistent with asthma prevalence being a near-static outcome with only 1.5% "
    "within-region variance in these data.")
new_c = ("After adjusting for region and year fixed effects, no robust association was "
    "identified, consistent with asthma prevalence being a near-static outcome with only 1.5% "
    "within-region variance in these data. This null/negative finding proved directionally "
    "stable but only narrowly robust: allowing each region its own secular trend, in addition "
    "to its own fixed level, shrank the coefficient by 63% and left it only narrowly "
    "significant, indicating that a meaningful share of the primary result may reflect "
    "differential regional trajectories rather than a shared within-region mechanism.")
full_c = "".join(r.text for r in conclusion_open.runs)
if old_c not in full_c:
    raise ValueError("Conclusion opening-paragraph anchor not found.")
conclusion_open.runs[0].text = full_c.replace(old_c, new_c)
for r in conclusion_open.runs[1:]:
    r.text = ""
print("Fix 2 applied: Conclusion opening paragraph aligned with the region-trends result.")

# =========================================================================
# 3. REFERENCES -- add Wolfers (2006), verified via web search:
#    Wolfers, J. (2006). Did Unilateral Divorce Laws Raise Divorce Rates?
#    A Reconciliation and New Results. American Economic Review, 96(5),
#    1802-1820.
# =========================================================================
ai_disclosure_h1 = find_para(text_exact="Artificial Intelligence Disclosure")
new_ref = ai_disclosure_h1.insert_paragraph_before("")
new_ref.paragraph_format.left_indent = Emu(273685)
new_ref.paragraph_format.first_line_indent = Emu(-273685)
new_ref.paragraph_format.space_after = Emu(88900)
style_run(new_ref.add_run(
    "28. Wolfers J. Did unilateral divorce laws raise divorce rates? A reconciliation and new "
    "results. American Economic Review. 2006;96(5):1802-1820. doi:10.1257/aer.96.5.1802"
), 11, bold=False)
print("Fix 3 applied: reference 28 (Wolfers, 2006) added.")

doc.save(PATH)

word_count_after = sum(len(p.text.split()) for p in doc.paragraphs)
print(f"\nWord count before (part 2): {word_count_before}")
print(f"Word count after (part 2):  {word_count_after}")
print(f"Net word change (part 2):   {word_count_after - word_count_before:+d}")
print(f"\nSaved {PATH}")
