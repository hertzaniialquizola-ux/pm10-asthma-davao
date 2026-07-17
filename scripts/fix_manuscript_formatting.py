"""
fix_manuscript_formatting.py
==============================
Two small post-hoc fixes to the edits made by update_manuscript.py, found
by rendering the docx to PDF and visually inspecting it:

1. The Limitations paragraph ended up with "Finally, ..." appearing twice
   (once in the pre-existing GBD-smoothing sentence, once in the newly
   appended MDE sentence). Retitles the newly appended sentence's lead-in
   to "In addition," to avoid the duplicate.

2. References 9-14 (added by update_manuscript.py) were missing the
   hanging-indent paragraph formatting that references 1-8 already have
   (left_indent=273685 EMU, first_line_indent=-273685 EMU, space_after=
   88900 EMU), so they visually wrapped flush-left instead of matching the
   existing reference-list style. This copies that exact formatting onto
   the new references.
"""

import docx
from docx.shared import Emu

PATH = "outputs/paper/PM25_Pediatric_Asthma_Philippines_REFORMATTED.docx"
doc = docx.Document(PATH)

# ---- Fix 1: duplicate "Finally" in Limitations ----
fixed_1 = False
for p in doc.paragraphs:
    if p.text.strip().startswith("Several limitations must be acknowledged"):
        for r in p.runs:
            if "Finally, a minimum-detectable-effect calculation" in r.text:
                r.text = r.text.replace(
                    "Finally, a minimum-detectable-effect calculation",
                    "In addition, a minimum-detectable-effect calculation",
                )
                fixed_1 = True
assert fixed_1, "Did not find the duplicate-Finally sentence to fix."
print("Fix 1 applied: de-duplicated 'Finally' in Limitations.")

# ---- Fix 2: hanging indent on references 9-14 ----
ref_prefixes = tuple(f"{n}. " for n in range(9, 15))
n_fixed = 0
for p in doc.paragraphs:
    t = p.text.strip()
    if t.startswith(ref_prefixes):
        pf = p.paragraph_format
        pf.left_indent = Emu(273685)
        pf.first_line_indent = Emu(-273685)
        pf.space_after = Emu(88900)
        n_fixed += 1
assert n_fixed == 6, f"Expected to fix 6 references, fixed {n_fixed}."
print(f"Fix 2 applied: hanging indent set on {n_fixed} references (9-14).")

doc.save(PATH)
print("Saved", PATH)
