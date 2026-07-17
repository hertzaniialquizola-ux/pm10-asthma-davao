"""
add_placebo_figure.py
=======================
Embeds Figure 10 (outputs/figures/placebo_test_comparison.png) with a
caption, right after the "Robustness Check: Placebo Outcome" body
paragraph added by update_manuscript_placebo.py (and therefore still
before "Discussion", continuing the figure-numbering sequence: 1-6
original, 7 jackknife, 8 lag/rolling-mean, 9 confounding DAG, 10 this one).
"""

import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

PATH = "outputs/paper/PM25_Pediatric_Asthma_Philippines_REFORMATTED.docx"
FONT = "Times New Roman"

doc = docx.Document(PATH)


def find_para(text_exact=None, text_startswith=None):
    for p in doc.paragraphs:
        t = p.text.strip()
        if text_exact is not None and t == text_exact:
            return p
        if text_startswith is not None and t.startswith(text_startswith):
            return p
    raise ValueError(f"paragraph not found: {text_exact or text_startswith}")


def style_run(run, size, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return run


def insert_figure(anchor, path, width_in=5.83):
    p = anchor.insert_paragraph_before("")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width_in))
    return p


def insert_caption(anchor, label, rest):
    p = anchor.insert_paragraph_before("")
    p.paragraph_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style_run(p.add_run(label), 10, bold=True, italic=True)
    style_run(p.add_run(rest), 10, bold=False, italic=True)
    return p


discussion_h1 = find_para(text_exact="Discussion")
insert_figure(discussion_h1, "outputs/figures/placebo_test_comparison.png")
insert_caption(discussion_h1, "Figure 10. ",
    "Placebo/negative-control test: two-way fixed-effects β for PM2.5 predicting asthma "
    "prevalence (primary outcome, left) versus low back pain prevalence (negative control, "
    "right; no plausible PM2.5 pathway). Points show β; bars show 95% confidence intervals "
    "(clustered SE). The negative control shows no significant association (p = 0.664), "
    "supporting the primary asthma null as a genuine finding."
)

doc.save(PATH)
print("Saved", PATH, "-- Figure 10 (placebo comparison) inserted.")
