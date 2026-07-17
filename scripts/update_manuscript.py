"""
update_manuscript.py
======================
Applies this session's robustness-upgrade write-up to
outputs/paper/PM25_Pediatric_Asthma_Philippines_REFORMATTED.docx (the
working manuscript file -- NOT outputs/paper/PM25_Pediatric_Asthma_STS_FORMATTED.docx
or any NHSJS-named file, which are separate submission variants and must
not be touched).

Matches the existing document's own formatting conventions (detected by
inspecting run/paragraph properties directly, since this document uses
direct formatting rather than named Word styles):
  H1 section heading : bold, 13pt, Times New Roman, space_before=16pt, space_after=8pt
  H2 subsection head  : bold, 11pt, Times New Roman, space_before=11pt, space_after=6pt
  Body paragraph      : regular, 11pt, Times New Roman, justified, space_after=8pt
  Figure/Table caption: "Label. " bold+italic 10pt, rest italic 10pt, justified, space_after=12pt

Run once; re-running will duplicate the inserted sections (there is no
idempotency guard, by design, so re-runs after further manual edits don't
silently no-op).
"""

import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

PATH = "outputs/paper/PM25_Pediatric_Asthma_Philippines_REFORMATTED.docx"
FONT = "Times New Roman"

doc = docx.Document(PATH)


def find_para(text_exact=None, text_startswith=None):
    for p in doc.paragraphs:
        t = p.text.strip()
        if text_exact is not None and t == text_exact:
            return p
        if text_startswith is not None and t.startswith(text_startswith):
            return p
    raise ValueError(f"paragraph not found: {text_exact or text_startswith}")


def style_run(run, size, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return run


def insert_h2(anchor, text):
    p = anchor.insert_paragraph_before("")
    p.paragraph_format.space_before = Pt(11)
    p.paragraph_format.space_after = Pt(6)
    style_run(p.add_run(text), 11, bold=True)
    return p


def insert_body(anchor, text):
    p = anchor.insert_paragraph_before("")
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style_run(p.add_run(text), 11, bold=False)
    return p


def insert_caption(anchor, label, rest):
    p = anchor.insert_paragraph_before("")
    p.paragraph_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style_run(p.add_run(label), 10, bold=True, italic=True)
    style_run(p.add_run(rest), 10, bold=False, italic=True)
    return p


def insert_figure(anchor, path, width_in=5.83):
    p = anchor.insert_paragraph_before("")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width_in))
    return p


# ─────────────────────────────────────────────────────────────────────────
# 1. METHODS — note the additional checks run this session
# ─────────────────────────────────────────────────────────────────────────
methods_para = find_para(text_startswith="where αᵢ is a region fixed effect")
methods_para.runs[-1].text += (
    " In a later session, five additional robustness checks were added to the primary "
    "coefficient — a wild cluster bootstrap, a leave-one-region-out jackknife, "
    "re-estimation excluding the 2020–2021 COVID-19 pandemic years, a Hausman "
    "specification test of fixed effects against random effects, and a Moran's I test "
    "for residual spatial autocorrelation — plus two exploratory analyses: lagged/"
    "cumulative PM2.5 exposure specifications and a minimum-detectable-effect "
    "calculation (see Results)."
)

# ─────────────────────────────────────────────────────────────────────────
# 2. RESULTS — insert 7 new subsections before "Discussion"
# ─────────────────────────────────────────────────────────────────────────
discussion_h1 = find_para(text_exact="Discussion")

insert_h2(discussion_h1, "Robustness Check: Wild Cluster Bootstrap")
insert_body(discussion_h1,
    "Because the label- and block-permutation tests above are one nonparametric approach "
    "to small-cluster inference, we additionally validated the primary coefficient with a "
    "second, methodologically distinct method: the wild cluster bootstrap, restricted "
    "version (WCR; Cameron, Gelbach & Miller, 2008), one of the most widely recommended "
    "small-G inference procedures in applied panel econometrics. Under the null hypothesis "
    "β = 0, we resampled the model’s own restricted residuals with region-level "
    "Rademacher sign flips (5,000 replications, seed = 42) and re-estimated the two-way "
    "fixed-effects coefficient each time, comparing the resulting bootstrap t-distribution "
    "to the actually observed t-statistic (t = −3.10). The WCR bootstrap p-value was "
    "p = 0.036 — still below the conventional 0.05 threshold, but a meaningfully weaker "
    "result than the clustered-SE p-value (p = 0.0024) or either permutation-test p-value "
    "(p = 0.0048 and p = 0.0008). This pattern — significance that survives but with "
    "much less margin under small-cluster-robust bootstrap inference — is consistent "
    "with the caution already noted in this manuscript about the modest number of clusters "
    "(17 regions) underlying the clustered standard errors."
)

insert_h2(discussion_h1, "Robustness Check: Leave-One-Region-Out Jackknife")
insert_body(discussion_h1,
    "To confirm the primary coefficient was not being driven by any single region — a "
    "natural concern given that the National Capital Region (NCR) is a clear outlier in "
    "levels (Figure 1; highest PM2.5 and highest asthma prevalence in the country) — we "
    "re-estimated the two-way fixed-effects model 17 times, each time dropping one region "
    "and refitting on the remaining 16 regions × 10 years (n = 160). The coefficient "
    "remained negative in all 17 refits (range: β = −1.416 to β = −2.812) "
    "and remained statistically significant (p < 0.05) in all 17 refits. The result was not "
    "uniform, however: dropping NCR specifically produced the smallest-magnitude, least "
    "precise estimate of the 17 (β = −1.416, SE = 0.699, p = 0.045) — a 45% "
    "reduction in magnitude relative to the full-sample β = −2.554, and the only "
    "refit with p above 0.03. Figure 7 shows all 17 estimates. NCR is therefore a "
    "meaningful, though not sole, contributor to the primary result’s magnitude and "
    "precision: the negative within-region association holds without NCR, but is markedly "
    "weaker."
)
insert_figure(discussion_h1, "outputs/figures/jackknife_leave_one_region_out.png")
insert_caption(discussion_h1, "Figure 7. ",
    "Leave-one-region-out jackknife: two-way fixed-effects β (PM2.5 → asthma "
    "prevalence) re-estimated 17 times, each time excluding one region, n = 160 per refit. "
    "Points show β with the labeled region excluded; horizontal bars show 95% "
    "confidence intervals (clustered SE). The dashed vertical line marks the full-sample "
    "β = −2.554 (n = 170). Excluding NCR (red point) produces the smallest-"
    "magnitude, least precise estimate of the 17."
)

insert_h2(discussion_h1, "Robustness Check: Excluding the COVID-19 Pandemic Years")
insert_body(discussion_h1,
    "The Limitations section below notes that the study period includes the COVID-19 "
    "pandemic years (2020–2021), which disrupted healthcare utilization and likely "
    "affected formal asthma diagnosis rates. To quantify that caveat, we re-estimated the "
    "primary model on a 17-region × 8-year panel excluding 2020 and 2021 (2013–2019 "
    "and 2022; n = 136). The coefficient attenuated from β = −2.554 (p = 0.0024, "
    "n = 170) to β = −2.113 (SE = 1.225, p = 0.088, n = 136) — a 17% reduction "
    "in magnitude, and a loss of statistical significance at the conventional 0.05 "
    "threshold (though not at 0.10). Removing two of ten years necessarily reduces "
    "statistical power (within-R² fell from 0.080 to 0.056), so this attenuation is "
    "consistent with reduced precision alone rather than a change in the underlying "
    "relationship; it cannot, however, rule out that pandemic-era diagnostic disruption "
    "contributed to the primary estimate, and this should be read as a genuine, unresolved "
    "sensitivity rather than a fully robust result."
)

insert_h2(discussion_h1, "Robustness Check: Lag Structure and Cumulative Exposure")
insert_body(discussion_h1,
    "The Discussion below argues that asthma prevalence, as a slowly changing “stock” "
    "measure, may be poorly suited to detecting a same-year PM2.5 effect, and that "
    "cumulative exposure is more biologically plausible than a single year’s mean "
    "concentration. We tested this directly by re-estimating the two-way fixed-effects "
    "model with PM2.5 lagged 1, 2, and 3 years, and with a 3-year trailing rolling-mean "
    "PM2.5 exposure, each on the correspondingly smaller (but still balanced) rectangular "
    "panel. The 1- and 2-year lag specifications produced similar-magnitude, still-negative "
    "coefficients (β = −2.156, p = 0.004, n = 153; β = −2.168, p = 0.054, "
    "n = 136); the 3-year lag attenuated further and lost significance (β = −1.627, "
    "p = 0.151, n = 119). The 3-year rolling-mean specification, by contrast, produced a "
    "substantially larger and more precisely estimated negative coefficient than the "
    "same-year model (β = −6.040, SE = 1.110, p < 0.0001, within-R² = 0.204, "
    "n = 136) — the strongest result, in either direction, of any specification "
    "examined in this manuscript. Figure 8 compares all five specifications. This is an "
    "important and unexpected finding that should be read carefully: a stronger negative "
    "within-region association under cumulative exposure does not support a protective "
    "effect of PM2.5, for the same reason stated throughout this manuscript — the "
    "negative sign reflects the absence of a positive signal in a dataset where asthma "
    "prevalence is nearly time-invariant within regions, not evidence that PM2.5 reduces "
    "asthma. What it does show is that the null/negative result is not an artifact "
    "specific to same-year exposure timing; if anything, a more biologically plausible "
    "cumulative-exposure specification strengthens the finding of no detectable positive "
    "within-region PM2.5–asthma association in this dataset."
)
insert_figure(discussion_h1, "outputs/figures/lag_structure_comparison.png")
insert_caption(discussion_h1, "Figure 8. ",
    "Two-way fixed-effects β under alternative PM2.5 exposure timing: same-year "
    "(primary spec), 1/2/3-year lags, and a 3-year trailing rolling mean. Points show "
    "β; horizontal bars show 95% confidence intervals (clustered SE). All "
    "specifications remain negative; the 3-year rolling mean produces the largest-"
    "magnitude, most precisely estimated coefficient of any specification in this "
    "manuscript (β = −6.040, p < 0.0001)."
)

insert_h2(discussion_h1, "Specification Test: Fixed Effects versus Random Effects (Hausman Test)")
insert_body(discussion_h1,
    "This manuscript’s Methods section states that the two-way fixed-effects model was "
    "chosen because it controls for stable regional differences, without a formal test of "
    "that choice against a random-effects (RE) alternative, which would be more efficient "
    "if region effects were uncorrelated with PM2.5 — an assumption this manuscript’s "
    "own Discussion argues is false (regions with structurally higher PM2.5 also have "
    "structurally higher urbanization and diagnostic capacity). We fit a Swamy-Arora "
    "random-effects estimator (Swamy & Arora, 1972) on the same regressor set (PM2.5 plus "
    "year dummies) and compared it to the fixed-effects estimate via a Hausman "
    "specification test (Hausman, 1978). The classical Hausman chi-square statistic was "
    "not reliably computable in this dataset: the estimated variance difference "
    "Var(β_FE) − Var(β_RE) was negative under both the two-way specification "
    "and a simpler one-way (region-only) specification, a known finite-sample degeneracy of "
    "the classical Hausman test that can occur when the random-effects estimator’s "
    "theoretical efficiency gain over fixed effects is small (here, the RE quasi-demeaning "
    "weight θ = 0.978, i.e. the RE transform is already close to a full within-"
    "transform because between-region variance in this dataset is very large relative to "
    "within-region variance). A valid Hausman p-value therefore cannot be reported. What "
    "can be reported directly: the RE point estimate (β = −2.126) differs from the "
    "FE point estimate (β = −2.554) by seventeen percent, moving toward the "
    "strongly positive pooled correlation (r = +0.887) that this manuscript’s central "
    "argument attributes to between-region confounding — exactly the direction RE "
    "contamination would be expected to move the estimate if that confounding is real. "
    "Combined with the qualitative argument already made in the Discussion, this supports "
    "retaining fixed effects as the primary specification even without a formally "
    "significant test statistic."
)

insert_h2(discussion_h1, "Statistical Power: Minimum Detectable Effect")
insert_body(discussion_h1,
    "To make the qualitative concern implicit in the low within-R² (= 0.080) explicit, "
    "we calculated the minimum detectable effect (MDE) implied by the primary model’s "
    "clustered standard error (SE = 0.825, two-way FE residual df ≈ 143). At "
    "conventional thresholds (α = 0.05, 80% power), this design could reliably detect "
    "a true |β| of at least 2.33 asthma cases per 100,000 per 1 µg/m³ PM2.5 "
    "— 91% of the actually observed |β| = 2.554. In other words, this study had "
    "approximately 80% power to detect the effect size it in fact found, but a true effect "
    "smaller than roughly 2.3 units could plausibly go undetected altogether. At a "
    "stricter 90%-power standard, the MDE (2.69) exceeds the observed coefficient, meaning "
    "this design is only borderline-adequately powered even for the effect actually "
    "estimated. This should be read alongside, not instead of, the within-R² = 0.080 "
    "already reported: both describe the same underlying limitation — a panel with "
    "very little within-region temporal variance in the outcome — from different "
    "angles."
)

insert_h2(discussion_h1, "Robustness Check: Spatial Autocorrelation of Residuals (Moran’s I)")
insert_body(discussion_h1,
    "The clustered standard errors used throughout this manuscript assume residuals are "
    "uncorrelated across regions (only within-region correlation over time is accounted "
    "for). If two-way fixed-effects residuals were spatially clustered — for example, "
    "if neighboring regions shared unmodeled pollution sources or health-system "
    "characteristics — that assumption would be violated and the reported standard "
    "errors could understate true uncertainty. We tested this using Moran’s I on the "
    "primary model’s residuals, averaged to one value per region. Because contiguity "
    "(“queen”/“rook”) adjacency is poorly suited to the Philippines’ "
    "archipelagic geography — many regions are island groups with no land border to "
    "any neighbor — spatial weights were instead built from real inter-region "
    "distances: province-level polygon centroids from the GADM shapefile already used "
    "elsewhere in this project for PM2.5 area-weighting were aggregated (area-weighted) to "
    "the 17 study regions, and a K-nearest-neighbor (k = 4) weight matrix was built from "
    "the resulting real great-circle distances. The observed Moran’s I was "
    "−0.0017, close to its expectation under no spatial autocorrelation (−0.063 "
    "for n = 17), and a spatial permutation test (5,000 reshuffles) gave p = 0.994 "
    "(two-sided). We find no evidence of residual spatial autocorrelation, supporting the "
    "validity of region-clustered (rather than spatially adjusted) standard errors for "
    "this dataset."
)

# ─────────────────────────────────────────────────────────────────────────
# 3. DISCUSSION — confounding DAG + E-value, then a sensitivity synthesis
# ─────────────────────────────────────────────────────────────────────────
why_wrong_outcome_h2 = find_para(text_exact="Why Asthma Prevalence Is the Wrong Outcome Variable")

insert_h2(why_wrong_outcome_h2, "Confounding Structure and an Approximate E-Value")
insert_figure(why_wrong_outcome_h2, "outputs/figures/fig7_confounding_dag.png")
insert_caption(why_wrong_outcome_h2, "Figure 9. ",
    "Directed acyclic graph of the hypothesized confounding structure. Blue: the path of "
    "interest (two-way fixed-effects estimate). Gray dashed: confounding paths closed by "
    "region and year fixed effects. Red: the one confounding path this design does not "
    "close (unobserved region × year-varying confounders), matching the Limitations "
    "section."
)
insert_body(why_wrong_outcome_h2,
    "Figure 9 lays out the argument made above as a directed acyclic graph (DAG): region "
    "(time-invariant urbanization, healthcare infrastructure, and diagnostic capacity) and "
    "year (shared national trend) both plausibly affect PM2.5 and asthma prevalence, and "
    "both paths are closed by the two-way fixed-effects design; an unresolved path — "
    "unobserved region × year confounders such as local healthcare-access changes, "
    "smoking prevalence, or socioeconomic shifts within a region over time (already named "
    "in Limitations) — remains open."
)
insert_body(why_wrong_outcome_h2,
    "As a complementary, approximate quantification of how strong an unmeasured "
    "confounder would need to be, we calculated an E-value (VanderWeele & Ding, 2017) for "
    "the pooled cross-sectional correlation (r = +0.887). Because the E-value framework is "
    "defined for risk ratios, not Pearson correlations, this required a chain of "
    "approximate conversions (r → Cohen’s d → odds ratio → risk ratio; "
    "Cohen, 1988) reported in full, with each step’s caveats, in outputs/tables/"
    "evalue_results.csv; the “rare outcome” step in particular is a poor fit "
    "here, since asthma prevalence in this dataset averages roughly 7.5%, not "
    "conventionally rare. With that important caveat, the resulting E-value is "
    "approximately 2,093 — an extremely large number, meaning an extremely strong "
    "confounder would be needed to produce the pooled association on its own. This should "
    "not be read as evidence for a causal PM2.5 effect. If anything, it is consistent with "
    "— not contrary to — this manuscript’s central argument: region-level "
    "confounding of exactly this magnitude and kind (urbanization, diagnostic capacity) is "
    "independently identified and adjusted for via fixed effects, which is exactly why the "
    "pooled association collapses under the two-way FE model. This E-value should be read "
    "as an approximation, not a precise bound, given the conversion chain above."
)

implications_h2 = find_para(text_exact="Implications for Philippine Environmental Health Policy")
insert_h2(implications_h2, "Synthesis of Sensitivity Analyses")
insert_body(implications_h2,
    "Across all robustness checks conducted for the primary coefficient — clustered "
    "standard errors, label- and block-permutation tests, a wild cluster bootstrap, a "
    "leave-one-region-out jackknife, and exclusion of the COVID-19 pandemic years — "
    "the two-way fixed-effects estimate was directionally consistent (always negative) and "
    "statistically significant in every check except one (excluding 2020–2021, where "
    "it narrowly missed conventional significance at p = 0.088). The estimate was "
    "measurably, though not fully, sensitive to two specific features of the data: the "
    "National Capital Region and the two pandemic years. Alternative exposure-timing "
    "specifications (1/2/3-year lags, 3-year rolling mean) did not overturn the negative "
    "sign at any lag, and a cumulative 3-year exposure measure produced the strongest "
    "result of any specification tested. A Hausman-style comparison against random effects "
    "could not be formally tested due to a known finite-sample degeneracy, but the "
    "substantial divergence between the fixed- and random-effects point estimates is "
    "itself informal evidence for the region-level confounding this manuscript’s "
    "central argument describes. Moran’s I found no evidence of residual spatial "
    "autocorrelation, supporting the region-clustered standard errors used throughout. "
    "Taken together, these checks support treating the primary null/negative finding as a "
    "genuine feature of this dataset rather than an artifact of any single modeling "
    "choice, while also making explicit exactly where that finding is least secure (NCR, "
    "the pandemic years, and the overall statistical power of the design, quantified above "
    "as an MDE of ≈ 2.33)."
)

# ─────────────────────────────────────────────────────────────────────────
# 4. LIMITATIONS — expand the population-weighting sentence, add MDE framing
# ─────────────────────────────────────────────────────────────────────────
limitations_para = find_para(text_startswith="Several limitations must be acknowledged")
old = limitations_para.runs[0].text
old = old.replace(
    "Province-to-region aggregation used simple means rather than population-weighted "
    "estimates, which may introduce aggregation bias. ",
    "Province-to-region aggregation used simple means rather than population-weighted "
    "estimates, which may introduce aggregation bias; a population-weighted aggregation "
    "was considered for this manuscript but not implemented, because the only population "
    "dataset available in this repository covers just 12 of the Philippines’ "
    "approximately 81 provinces and only through 2020 (missing two of this study’s ten "
    "years), and was judged insufficient to redo the aggregation credibly for the full "
    "panel. "
)
limitations_para.runs[0].text = old
limitations_para.runs[0].text += (
    " Finally, a minimum-detectable-effect calculation (see Results) indicates this design "
    "could reliably detect true effects only down to roughly 90% of the magnitude actually "
    "observed at conventional 80% power — a concrete bound on, rather than merely a "
    "qualitative caveat about, the study’s statistical power."
)

# ─────────────────────────────────────────────────────────────────────────
# 5. ABSTRACT — append sensitivity summary to Results and Conclusions
# ─────────────────────────────────────────────────────────────────────────
abstract_para = find_para(text_startswith="Background.")
runs = abstract_para.runs
# runs[7] = Results body, runs[9] = Conclusions body (see update_manuscript.py docstring)
results_run = None
conclusions_run = None
for i, r in enumerate(runs):
    if r.text.strip() == "Results.":
        results_run = runs[i + 1]
    if r.text.strip() == "Conclusions.":
        conclusions_run = runs[i + 1]
assert results_run is not None and conclusions_run is not None

results_run.text += (
    "Sensitivity analyses were mostly, but not uniformly, supportive: a wild cluster "
    "bootstrap (p = 0.036) and a leave-one-region-out jackknife (β ranged "
    "−1.416 to −2.812 across 17 refits, all p < 0.05) preserved the "
    "coefficient’s sign and significance, while excluding the COVID-19 pandemic years "
    "(2020–2021) attenuated it to non-significance (β = −2.113, p = 0.088) "
    "and a 3-year cumulative PM2.5 exposure specification instead produced a substantially "
    "larger, highly significant negative coefficient (β = −6.040, p < 0.0001). "
)
conclusions_run.text += (
    " These sensitivity analyses show the fixed-effects estimate is reasonably, though not "
    "perfectly, robust: it is sensitive to the National Capital Region and to the "
    "COVID-19 pandemic years specifically, and strengthens rather than weakens under a "
    "more biologically plausible cumulative-exposure specification."
)

# ─────────────────────────────────────────────────────────────────────────
# 6. REFERENCES — add citations for newly cited methods (flagged for the
#    author's own verification in the session status doc, same convention
#    already used for references 1-5 in this repo's history)
# ─────────────────────────────────────────────────────────────────────────
ai_disclosure_h1 = find_para(text_exact="Artificial Intelligence Disclosure")
new_refs = [
    "9. Cameron AC, Gelbach JB, Miller DL. Bootstrap-based improvements for inference "
    "with clustered errors. The Review of Economics and Statistics. 2008;90(3):414-427.",
    "10. Hausman JA. Specification tests in econometrics. Econometrica. 1978;46(6):1251-1271.",
    "11. Swamy PAVB, Arora SS. The exact finite sample properties of the estimators of "
    "coefficients in the error components regression models. Econometrica. 1972;40(2):261-275.",
    "12. VanderWeele TJ, Ding P. Sensitivity analysis in observational research: "
    "introducing the E-value. Annals of Internal Medicine. 2017;167(4):268-274.",
    "13. Cohen J. Statistical Power Analysis for the Behavioral Sciences. 2nd ed. "
    "Hillsdale, NJ: Lawrence Erlbaum Associates; 1988.",
    "14. Wooldridge JM. Econometric Analysis of Cross Section and Panel Data. 2nd ed. "
    "Cambridge, MA: MIT Press; 2010.",
]
for ref_text in new_refs:
    p = ai_disclosure_h1.insert_paragraph_before("")
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style_run(p.add_run(ref_text), 11, bold=False)
    # NOTE: insert_paragraph_before(ai_disclosure_h1) inserts directly before
    # that heading (i.e. right after reference 8, the last existing
    # reference), so calling this repeatedly with the same anchor builds up
    # the inserted paragraphs in the order this loop runs them -- ref 9 ends
    # up right after ref 8, ref 10 right after ref 9, etc. (same pattern used
    # for all the Results/Discussion insertions above).

doc.save(PATH)
print("Saved", PATH)
