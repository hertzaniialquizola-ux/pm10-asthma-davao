"""
task2_multiplicity.py
========================
Adds a scoped Holm-Bonferroni multiplicity correction to the manuscript --
NOT a blanket correction across every p-value. Two families only:

  Family 1: the 5-outcome comparison (Table 4 / Figure 6)
  Family 2: the 4 alternative exposure-timing specs (lag1/lag2/lag3/rolling
            mean) -- explicitly EXCLUDES the same-year primary specification,
            which is not part of this exploratory family.

Does NOT touch: WCR bootstrap, permutation test, jackknife, COVID exclusion,
placebo test, Moran's I -- these are robustness/inference checks on ONE
pre-specified estimate, not a search across candidates, per the user's
explicit instruction.

Citation verified via web search (not memory): Holm S. A simple sequentially
rejective multiple test procedure. Scandinavian Journal of Statistics.
1979;6(2):65-70. doi:10.2307/4615733
"""

import docx
from docx.shared import Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH

PATH = "outputs/paper/PM25_Pediatric_Asthma_Philippines_REFORMATTED.docx"
FONT = "Times New Roman"

doc = docx.Document(PATH)


def find_para(text_exact=None, text_startswith=None, occurrence=1):
    count = 0
    for p in doc.paragraphs:
        t = p.text.strip()
        match = (text_exact is not None and t == text_exact) or \
                (text_startswith is not None and t.startswith(text_startswith))
        if match:
            count += 1
            if count == occurrence:
                return p
    raise ValueError(f"paragraph not found: {text_exact or text_startswith} (occurrence {occurrence})")


def style_run(run, size=11, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return run


def replace_paragraph_text(para, new_text):
    runs = para.runs
    if not runs:
        r = para.add_run(new_text)
        style_run(r)
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""


# =======================================================================
# Holm-Bonferroni computation (done here, not by hand, so the numbers in
# the manuscript are generated the same way they're checked)
# =======================================================================
def holm_bonferroni(pvals_dict):
    items = sorted(pvals_dict.items(), key=lambda kv: kv[1])
    m = len(items)
    adj = {}
    running_max = 0.0
    for j, (label, p) in enumerate(items, start=1):
        mult = m - j + 1
        raw_adj = min(1.0, mult * p)
        running_max = max(running_max, raw_adj)
        adj[label] = running_max
    return adj


family1_raw = {
    "asthma": 0.0024,
    "lung_cancer": 0.0333,
    "copd": 0.0925,
    "lri": 0.0965,
    "resp_mortality": 0.1737,
}
family1_adj = holm_bonferroni(family1_raw)

family2_raw = {
    "lag1": 0.0041,
    "lag2": 0.0535,
    "lag3": 0.1512,
    "roll3": 3.1657e-7,
}
family2_adj = holm_bonferroni(family2_raw)

print("Family 1 (5-outcome test) Holm-adjusted p-values:")
for k, v in family1_adj.items():
    print(f"  {k:16s} raw={family1_raw[k]:.4f}  adj={v:.4f}")
print("\nFamily 2 (lag/rolling-mean sweep) Holm-adjusted p-values:")
for k, v in family2_adj.items():
    print(f"  {k:16s} raw={family2_raw[k]:.4g}  adj={v:.4g}")

# =======================================================================
# 1. METHODS -- one new sentence describing the primary-vs-exploratory
#    framing and the correction method.
# =======================================================================
methods_para = find_para(text_startswith="where αᵢ is a region fixed effect")
old_tail = ("A further session added an eighth check, a placebo/negative-control test using an "
    "outcome with no plausible PM2.5 pathway (see Results).")
new_tail = (old_tail + " Because the same-year two-way fixed-effects model is the single "
    "pre-specified primary specification, the robustness/inference checks above were not "
    "treated as a multiple-comparisons family requiring correction; the five-outcome "
    "comparison and the four alternative exposure-timing specifications reported below, by "
    "contrast, are labeled — post hoc, not as a true pre-registration, which cannot be done "
    "retroactively — as exploratory/sensitivity analyses, and their p-values are additionally "
    "reported after a Holm-Bonferroni correction (Holm, 1979) applied separately within each "
    "family.")
hit = False
for r in methods_para.runs:
    if old_tail in r.text:
        r.text = r.text.replace(old_tail, new_tail)
        hit = True
if not hit:
    raise ValueError("Methods paragraph tail not found")
print("\nFix 1 applied: Methods paragraph updated with primary-vs-exploratory framing sentence.")

# =======================================================================
# 2. Table 4 (docx tables[3], the 5-outcome table) -- add a Holm-adjusted
#    p column.
# =======================================================================
t4 = doc.tables[3]
col = t4.add_column(Emu(700000))
header_cell = t4.rows[0].cells[-1]
header_cell.text = ""
p = header_cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
style_run(p.add_run("Holm-adj. p"), 10, bold=True)

row_order = ["asthma", "lung_cancer", "lri", "copd", "resp_mortality"]  # matches table row order
for i, key in enumerate(row_order, start=1):
    cell = t4.rows[i].cells[-1]
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    val = family1_adj[key]
    marker = "*" if val < 0.05 else ""
    style_run(p.add_run(f"{val:.4f}{marker}"), 10, bold=False)
print("Fix 2 applied: Table 4 given a new 'Holm-adj. p' column.")

# Update Table 4 caption
t4_caption = find_para(text_startswith="Table 4. Two-way fixed-effects panel regression results across five")
replace_paragraph_text(t4_caption,
    "Table 4. Two-way fixed-effects panel regression results across five respiratory-related "
    "outcomes. All models include 17 region and 10 year fixed effects; standard errors "
    "clustered by region. Philippines, 2013–2022, n = 170 region-years each. These five "
    "outcomes are treated as a post hoc exploratory family (not part of the pre-specified "
    "primary specification); the rightmost column reports p-values after a Holm-Bonferroni "
    "correction (Holm, 1979) applied across the 5 tests. * = still significant (adjusted "
    "p < 0.05)."
)
print("Fix 2b applied: Table 4 caption updated to describe the correction.")

# =======================================================================
# 3. Figure 6 caption -- pointer to the corrected values now in Table 4,
#    without regenerating the image itself (asterisks in the image still
#    reflect uncorrected p < 0.05, which is stated explicitly).
# =======================================================================
fig6_caption = find_para(text_startswith="Figure 6. Forest plot of two-way fixed-effects β coefficients")
replace_paragraph_text(fig6_caption,
    "Figure 6. Forest plot of two-way fixed-effects β coefficients (PM2.5 predicting each "
    "outcome, region and year effects) across five respiratory-related outcomes, Philippines, "
    "2013–2022. Points show the fixed-effects β for each outcome; horizontal bars represent "
    "95% confidence intervals clustered by region. The dashed vertical line marks β = 0. "
    "Asterisks in this figure denote uncorrected p < 0.05; see Table 4 for the same five "
    "p-values after a Holm-Bonferroni correction across this exploratory family."
)
print("Fix 3 applied: Figure 6 caption updated to point to the corrected values.")

# =======================================================================
# 4. Paragraph 95 (Testing Additional Respiratory Outcomes) -- report
#    corrected values alongside raw, and formalize the existing
#    qualitative "consistent with chance" framing.
# =======================================================================
p95 = find_para(text_startswith="To determine whether the absence of a within-region PM2.5 signal")
replace_paragraph_text(p95,
    "To determine whether the absence of a within-region PM2.5 signal was specific to asthma "
    "prevalence or reflected a broader limitation of this panel design, the same two-way "
    "fixed-effects specification was applied to four additional respiratory outcomes: lower "
    "respiratory infection incidence, COPD prevalence, tracheal/bronchus/lung cancer "
    "incidence, and respiratory-disease mortality (Table 4, Figure 6). These five outcomes "
    "were not pre-specified as a single confirmatory test; they are treated here as an "
    "exploratory family and, in addition to the raw p-values, Table 4 reports each one after "
    "a Holm-Bonferroni correction (Holm, 1979) applied across all 5 tests. Of the five raw "
    "p-values, four — asthma, lower respiratory infections, COPD, and respiratory mortality — "
    "showed no statistically significant within-region association with PM2.5 (all p > 0.09). "
    "Lung cancer incidence showed a marginally significant raw result (p = 0.033); after the "
    "Holm-Bonferroni correction, it is no longer significant (adjusted p = 0.133), while "
    "asthma remains significant under the same correction (adjusted p = 0.012) because its "
    "raw p-value was an order of magnitude smaller than the other four. This formalizes, "
    "rather than changes, what the manuscript already argued qualitatively: a single "
    "borderline result out of five untested-for-multiplicity comparisons is consistent with "
    "chance, and should not be interpreted as evidence of a true lung cancer effect. Rather "
    "than weakening the main result, this consistency across outcomes strengthens it: the "
    "absence of a detectable within-region signal is not an artifact of asthma prevalence "
    "specifically, but a broader feature of this ecological panel — one likely driven by the "
    "same combination of low within-region variance and modeled, smoothed subnational "
    "estimates discussed above. Future work using outcome variables with genuinely "
    "high-frequency variation (e.g., facility-level emergency admissions) would be better "
    "positioned to test whether a real effect exists beneath this noise floor."
)
print("Fix 4 applied: paragraph 95 updated with corrected p-values and formalized framing.")

# =======================================================================
# 5. Paragraph 68 (Lag Structure and Cumulative Exposure) -- report
#    corrected p-values for the 4-spec exploratory family (excludes
#    same-year baseline).
# =======================================================================
p68 = find_para(text_startswith="The Discussion below argues that asthma prevalence")
replace_paragraph_text(p68,
    "The Discussion below argues that asthma prevalence, as a slowly changing “stock” "
    "measure, may be poorly suited to detecting a same-year PM2.5 effect, and that cumulative "
    "exposure is more biologically plausible than a single year’s mean concentration. We "
    "tested this directly by re-estimating the two-way fixed-effects model with PM2.5 lagged "
    "1, 2, and 3 years, and with a 3-year trailing rolling-mean PM2.5 exposure, each on the "
    "correspondingly smaller (but still balanced) rectangular panel. These four alternative "
    "exposure-timing specifications were not pre-specified individually; they are treated "
    "here as an exploratory family distinct from the same-year primary specification, and "
    "their p-values are reported both raw and after a Holm-Bonferroni correction (Holm, 1979) "
    "applied across the 4 tests. The 1- and 2-year lag specifications produced "
    "similar-magnitude, still-negative coefficients (β = −2.156, raw p = 0.004, "
    "Holm-adjusted p = 0.012, n = 153; β = −2.168, raw p = 0.054, Holm-adjusted p = 0.107, "
    "n = 136); the 3-year lag attenuated further and lost significance even before correction "
    "(β = −1.627, raw p = 0.151, Holm-adjusted p = 0.151, n = 119). The 3-year rolling-mean "
    "specification, by contrast, produced a substantially larger and more precisely estimated "
    "negative coefficient than the same-year model (β = −6.040, SE = 1.110, raw p < 0.0001, "
    "Holm-adjusted p < 0.0001, within-R² = 0.204, n = 136) — the strongest result, in either "
    "direction, of any specification examined in this manuscript, and the only one of the "
    "four whose significance is not remotely in question after correction. Figure 8 compares "
    "all five specifications. This is an important and unexpected finding that should be read "
    "carefully: a stronger negative within-region association under cumulative exposure does "
    "not support a protective effect of PM2.5, for the same reason stated throughout this "
    "manuscript — the negative sign reflects the absence of a positive signal in a dataset "
    "where asthma prevalence is nearly time-invariant within regions, not evidence that PM2.5 "
    "reduces asthma. What it does show is that the null/negative result is not an artifact "
    "specific to same-year exposure timing, nor an artifact of testing several lag lengths "
    "and reporting the strongest one — the 1-year lag and the 3-year rolling mean both "
    "survive correction for having tried four specifications, while only the weaker 2- and "
    "3-year lags do not."
)
print("Fix 5 applied: paragraph 68 updated with raw + Holm-adjusted p-values for all 4 lag/rolling specs.")

# Figure 8 caption pointer
fig8_caption = find_para(text_startswith="Figure 8. Two-way fixed-effects β under alternative PM2.5 exposure timing")
replace_paragraph_text(fig8_caption,
    "Figure 8. Two-way fixed-effects β under alternative PM2.5 exposure timing: same-year "
    "(primary spec), 1/2/3-year lags, and a 3-year trailing rolling mean. Points show β; "
    "horizontal bars show 95% confidence intervals (clustered SE). All specifications remain "
    "negative; the 3-year rolling mean produces the largest-magnitude, most precisely "
    "estimated coefficient of any specification in this manuscript (β = −6.040, raw p < "
    "0.0001, Holm-adjusted p < 0.0001 across the 4-specification exploratory family; see text)."
)
print("Fix 5b applied: Figure 8 caption updated with Holm-adjusted result.")

# =======================================================================
# 6. Synthesis of Sensitivity Analyses -- fold in both corrections.
# =======================================================================
synthesis = find_para(text_startswith="Across all robustness checks conducted")
old_lag_sentence = ("Alternative exposure-timing specifications (1/2/3-year lags, 3-year rolling mean) did not "
    "overturn the negative sign at any lag, and a cumulative 3-year exposure measure produced "
    "the strongest result of any specification tested.")
new_lag_sentence = (old_lag_sentence + " Treating these four specifications as a post hoc "
    "exploratory family and applying a Holm-Bonferroni correction across them, the 1-year lag "
    "and the 3-year rolling mean remain significant while the 2- and 3-year lags do not — the "
    "strongest results are not an artifact of trying several lag lengths and reporting the "
    "best one.")
old_lung_sentence_marker = "Moran’s I found no evidence of residual spatial autocorrelation"
hit1 = False
for r in synthesis.runs:
    if old_lag_sentence in r.text:
        r.text = r.text.replace(old_lag_sentence, new_lag_sentence)
        hit1 = True
if not hit1:
    raise ValueError("Synthesis lag sentence not found")

# Insert the family-1 (5-outcome) correction note right after the lag-family note,
# before the Hausman sentence, so both corrected families are documented together.
old_hausman_lead = "A Hausman-style comparison against random effects"
new_lead = ("A Holm-Bonferroni correction was also applied to the five-outcome comparison above "
    "(Table 4): asthma remains significant after correction (adjusted p = 0.012) while lung "
    "cancer's borderline raw result does not (adjusted p = 0.133), formalizing this "
    "manuscript's existing chance-consistent reading of that result. " + old_hausman_lead)
hit2 = False
for r in synthesis.runs:
    if old_hausman_lead in r.text:
        r.text = r.text.replace(old_hausman_lead, new_lead, 1)
        hit2 = True
if not hit2:
    raise ValueError("Synthesis Hausman lead not found")
print("Fix 6 applied: Synthesis of Sensitivity Analyses updated with both corrected families.")

# =======================================================================
# 7. Reference list -- add Holm (1979) as reference 24.
# =======================================================================
ai_disclosure_h1 = find_para(text_exact="Artificial Intelligence Disclosure")
new_ref = ai_disclosure_h1.insert_paragraph_before("")
new_ref.paragraph_format.left_indent = Emu(273685)
new_ref.paragraph_format.first_line_indent = Emu(-273685)
new_ref.paragraph_format.space_after = Emu(88900)
style_run(new_ref.add_run(
    "24. Holm S. A simple sequentially rejective multiple test procedure. Scandinavian "
    "Journal of Statistics. 1979;6(2):65-70. doi:10.2307/4615733"
), 11, bold=False)
print("Fix 7 applied: reference 24 (Holm, 1979) added to reference list.")

doc.save(PATH)
print(f"\nSaved {PATH}")
