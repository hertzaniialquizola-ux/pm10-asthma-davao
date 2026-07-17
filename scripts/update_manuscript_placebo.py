"""
update_manuscript_placebo.py
==============================
Adds the 8th and final robustness check (placebo/negative-control outcome,
"Low back pain") to the manuscript: a new Results subsection (matching the
naming/style/detail of the other 7), a Methods sentence, and an update to
the existing "Synthesis of Sensitivity Analyses" Discussion subsection.
Reference 15 (Lipsitch, Tchetgen Tchetgen & Cohen, 2010) already exists in
the reference list as an intentional orphan (added in a prior session in
anticipation of this check) -- this script gives it its first real in-text
citation; no reference-list changes needed.
"""

import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

PATH = "outputs/paper/PM25_Pediatric_Asthma_Philippines_REFORMATTED.docx"
FONT = "Times New Roman"

doc = docx.Document(PATH)


def find_para(text_exact=None, text_startswith=None, text_contains=None):
    for p in doc.paragraphs:
        t = p.text.strip()
        if text_exact is not None and t == text_exact:
            return p
        if text_startswith is not None and t.startswith(text_startswith):
            return p
        if text_contains is not None and text_contains in t:
            return p
    raise ValueError(f"paragraph not found: {text_exact or text_startswith or text_contains}")


def style_run(run, size, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return run


def insert_h2(anchor, text):
    p = anchor.insert_paragraph_before("")
    p.paragraph_format.space_before = Pt(11)
    p.paragraph_format.space_after = Pt(6)
    style_run(p.add_run(text), 11, bold=True)
    return p


def insert_body(anchor, text):
    p = anchor.insert_paragraph_before("")
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style_run(p.add_run(text), 11, bold=False)
    return p


def replace_in_runs(para, old, new, required=True):
    hit = False
    for r in para.runs:
        if old in r.text:
            r.text = r.text.replace(old, new)
            hit = True
    if required and not hit:
        raise ValueError(f"text not found in paragraph: {old!r}")
    return hit


# ─────────────────────────────────────────────────────────────────────────
# 1. NEW RESULTS SUBSECTION — inserted right before "Discussion" (i.e. right
#    after the Moran's I subsection, the last of the existing 7)
# ─────────────────────────────────────────────────────────────────────────
discussion_h1 = find_para(text_exact="Discussion")

insert_h2(discussion_h1, "Robustness Check: Placebo Outcome")
insert_body(discussion_h1,
    "The multi-outcome analysis above (Table 4, Figure 6) tests generalizability across four "
    "additional respiratory-related outcomes, but all five outcomes examined so far (asthma "
    "plus the four in Table 4) share at least a plausible PM2.5 pathway. A stronger test of "
    "whether the primary null reflects a real absence of a same-year effect — rather than an "
    "artifact of how GBD subnational estimates are constructed, or of the two-way "
    "fixed-effects estimator itself — is a true negative control: an outcome with no "
    "plausible PM2.5 pathway at all (Lipsitch, Tchetgen Tchetgen & Cohen, 2010). We obtained "
    "GBD 2023 prevalence data for low back pain (ages 5–14, the same subnational "
    "province-level export filters used for the other outcomes) and confirmed non-degenerate "
    "variance across provinces and years before proceeding: mean prevalence was 773.0 ± 14.7 "
    "per 100,000, no province-year value was zero or duplicated, and every one of the 82 "
    "provinces showed genuine year-to-year variation (within-province SD ranging 3.3 to 14.5), "
    "ruling out a flat or degenerate series. Applying the identical two-way fixed-effects "
    "specification (region and year effects, standard errors clustered by region) used "
    "throughout this manuscript, PM2.5 showed no statistically significant association with "
    "low back pain prevalence (β = −0.180, SE = 0.414, p = 0.664, within-R² = 0.003, "
    "n = 170). This null result on a genuinely implausible outcome supports interpreting the "
    "primary asthma null as a real absence of a detectable same-year, within-region PM2.5 "
    "effect, rather than as an artifact of the GBD subnational estimation pipeline or of the "
    "fixed-effects estimator producing spurious associations generally."
)
print("Fix 1 applied: new 'Robustness Check: Placebo Outcome' subsection inserted before Discussion.")

# ─────────────────────────────────────────────────────────────────────────
# 2. METHODS — one new sentence noting the 8th check
# ─────────────────────────────────────────────────────────────────────────
methods_para = find_para(text_startswith="where αᵢ is a region fixed effect")
replace_in_runs(
    methods_para,
    "a minimum-detectable-effect calculation (see Results).",
    "a minimum-detectable-effect calculation (see Results). A further session added an eighth "
    "check, a placebo/negative-control test using an outcome with no plausible PM2.5 pathway "
    "(see Results).",
)
print("Fix 2 applied: Methods paragraph updated with one new sentence.")

# ─────────────────────────────────────────────────────────────────────────
# 3. SYNTHESIS OF SENSITIVITY ANALYSES — fold in the placebo result
# ─────────────────────────────────────────────────────────────────────────
synthesis_para = find_para(text_startswith="Across all robustness checks conducted")
replace_in_runs(
    synthesis_para,
    "Moran’s I found no evidence of residual spatial autocorrelation, supporting the "
    "region-clustered standard errors used throughout. Taken together,",
    "Moran’s I found no evidence of residual spatial autocorrelation, supporting the "
    "region-clustered standard errors used throughout. A placebo/negative-control test using "
    "an outcome with no plausible PM2.5 pathway (low back pain prevalence) found no "
    "significant association (p = 0.664), supporting the primary null as a genuine finding "
    "rather than an artifact of the GBD estimation pipeline. Taken together,",
)
print("Fix 3 applied: Synthesis of Sensitivity Analyses updated with the placebo result.")

doc.save(PATH)
print("\nSaved", PATH)
