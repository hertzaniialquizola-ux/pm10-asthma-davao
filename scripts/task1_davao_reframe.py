"""
task1_davao_reframe.py
========================
Gives the ACAG V6GL03-vs-V6GL02.04 satellite product-version-drift finding
(currently buried inside the Davao ground-truth validation subsection) its
own subsection identity, with explicit framing of why it matters beyond
this paper. No new analysis or data -- restructuring/reframing only, using
numbers already in the manuscript (MAE 1.48 vs 3.65 ug/m3, +2.87 ug/m3
average shift).

Net page-count effect is tracked, not assumed: bullets 109-112 in "Data"
and paragraphs 122/123 in "Interpretation" are tightened to make room for
the new subsection heading + framing paragraph.
"""

import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

PATH = "outputs/paper/PM25_Pediatric_Asthma_Philippines_REFORMATTED.docx"
FONT = "Times New Roman"

doc = docx.Document(PATH)


def find_para(text_exact=None, text_startswith=None, occurrence=1):
    count = 0
    for p in doc.paragraphs:
        t = p.text.strip()
        match = (text_exact is not None and t == text_exact) or \
                (text_startswith is not None and t.startswith(text_startswith))
        if match:
            count += 1
            if count == occurrence:
                return p
    raise ValueError(f"paragraph not found: {text_exact or text_startswith} (occurrence {occurrence})")


def style_run(run, size=11, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return run


def insert_h2(anchor, text):
    p = anchor.insert_paragraph_before("")
    p.paragraph_format.space_before = Pt(11)
    p.paragraph_format.space_after = Pt(6)
    style_run(p.add_run(text), 11, bold=True)
    return p


def insert_body(anchor, text):
    p = anchor.insert_paragraph_before("")
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style_run(p.add_run(text), 11, bold=False)
    return p


def replace_paragraph_text(para, new_text):
    """Replace a paragraph's visible text in-place, preserving its first
    run's formatting (and hence bullet/list formatting on the paragraph
    itself, which lives on the paragraph, not the run)."""
    runs = para.runs
    if not runs:
        r = para.add_run(new_text)
        style_run(r)
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""


def delete_paragraph(para):
    p = para._element
    p.getparent().remove(p)


word_count_before = sum(len(p.text.split()) for p in doc.paragraphs)

# ---------------------------------------------------------------------
# 1. Trim bullets 109-112 ("Data" caveats) to make room -- same facts,
#    tighter wording.
# ---------------------------------------------------------------------
b1 = find_para(text_startswith="Station 07/17 continuity.")
replace_paragraph_text(b1,
    "Station 07/17 continuity. EMB relocated this station (J.P. Laurel Avenue to Marfori "
    "Heights) around 2023; its reports disagree on the exact date. The retained series uses "
    "the 2021–2022 values both reports agree on, but may span two distinct physical locations "
    "rather than one continuously monitored site."
)

b2 = find_para(text_startswith="2020 data-capture flag.")
replace_paragraph_text(b2,
    "2020 data-capture flag. EMB’s 2022 report notes Stations 15 and 16 fell below the "
    "required 75% annual data-capture rate in 2020; the reported values (26.42, 24.48 µg/m³ "
    "PM10) are the only figures available and should be treated as less reliable than other "
    "years."
)

b3 = find_para(text_startswith="2024 precision.")
replace_paragraph_text(b3,
    "2024 precision. 2024 values are rounded whole numbers from a three-year summary table "
    "in the 2025 RSOBER, unlike the decimal-precision figures reported for 2020–2023 in the "
    "earlier reports."
)

b4 = find_para(text_startswith="2024 data substitution.")
replace_paragraph_text(b4,
    "2024 data substitution. The 2025 RSOBER notes May 2024 manual-station PM10 readings "
    "involved EMB-approved data substitution to meet minimum coverage requirements — this "
    "does not change the reported annual mean but is disclosed for provenance."
)
print("Fix 1 applied: Data caveat bullets (109-112) tightened.")

# ---------------------------------------------------------------------
# 2. Remove the "Version sensitivity check." bullet from Data -- its
#    method description is folded into the new subsection below instead
#    of being stated twice.
# ---------------------------------------------------------------------
vsc = find_para(text_startswith="Version sensitivity check.")
delete_paragraph(vsc)
print("Fix 2 applied: 'Version sensitivity check.' bullet removed from Data (content moved).")

# ---------------------------------------------------------------------
# 3. New H2 subsection, inserted right before the "Table 6." caption --
#    this splits what was one undifferentiated "Results" block into
#    "Results" (Table 5, the core ground-truth comparison) and this new,
#    separately-headed subsection (Table 6, the version-drift finding).
# ---------------------------------------------------------------------
table6_caption = find_para(text_startswith="Table 6. Version sensitivity check")
insert_h2(table6_caption, "Satellite Product-Version Drift: A Caution Beyond This Study")
insert_body(table6_caption,
    "Because the main regional panel uses an earlier ACAG release (V6.GL.02.04) while the "
    "Davao comparison above used the current release (V6GL03), the same Davao City clip was "
    "re-run on the three overlapping years (2020–2022) using the older V6.GL.02.04 files, to "
    "test whether the satellite-ground gap reflected atmospheric reality or a product-version "
    "artifact (Table 6)."
)
print("Fix 3 applied: new H2 'Satellite Product-Version Drift' subsection heading + intro inserted.")

# ---------------------------------------------------------------------
# 4. New generalization paragraph, inserted after the existing finding
#    paragraph (MAE 1.48 vs 3.65) and before "Interpretation" -- this is
#    the actual "why this matters beyond this paper" framing the task
#    asked for.
# ---------------------------------------------------------------------
interpretation_h2 = find_para(text_exact="Interpretation")
insert_body(interpretation_h2,
    "This is not a Davao-specific or single-paper concern. Any ecological PM2.5 epidemiology "
    "study that mixes satellite product releases across time — or upgrades mid-study to a "
    "newer release, as most groups extending a panel forward in time eventually will — risks "
    "introducing a spurious trend, or masking a real one, purely from a change in the "
    "underlying retrieval algorithm rather than a change in atmospheric PM2.5. Researchers "
    "using ACAG or comparable modeled satellite PM2.5 products for multi-year subnational "
    "panels should run a within-location, cross-version check of this kind before treating a "
    "multi-year satellite series as internally consistent, particularly around known "
    "release-transition years. The main 17-region panel in this study avoids this specific "
    "risk because it uses a single ACAG release (V6.GL.02.04) consistently across all 17 "
    "regions and all 10 years; this check is reported here as a secondary methodological "
    "contribution of this work, not merely as a footnote to the Davao ground-truth comparison "
    "above."
)
print("Fix 4 applied: generalization/caution paragraph added.")

# ---------------------------------------------------------------------
# 5. Trim "Interpretation" para 122 -- its first conclusion (the version-
#    drift point) is now stated, at greater length, in the new subsection
#    above; keep only the second conclusion (internal validity), and
#    trim it.
# ---------------------------------------------------------------------
p122 = find_para(text_startswith="This validation exercise supports two conclusions")
replace_paragraph_text(p122,
    "Together with the version check above, this validation exercise supports one further "
    "conclusion relevant to the main study: because the main 17-region panel uses a single "
    "ACAG release (V6.GL.02.04) consistently across all years and regions, the "
    "version-sensitivity issue documented above does not affect its internal validity — it "
    "matters chiefly for future work that mixes ACAG releases across time or extends the "
    "panel using newer satellite products."
)
print("Fix 5 applied: Interpretation paragraph 122 trimmed (removed now-redundant first conclusion).")

# ---------------------------------------------------------------------
# 6. Light trim of paragraph 123 (residual gap discussion) -- same
#    content, tighter wording.
# ---------------------------------------------------------------------
p123 = find_para(text_startswith="The residual gap that persists")
replace_paragraph_text(p123,
    "The residual gap that persists even after accounting for the version effect (most "
    "visibly in 2020–2021) likely reflects some combination of spatial-averaging differences "
    "between a 0.1° satellite pixel and point-source ground monitors, the ground stations’ "
    "documented 2020 data-capture shortfall, and genuine limits on satellite retrieval "
    "accuracy over a coastal, cloud-prone tropical city — this study cannot distinguish "
    "between them. Each would also apply, to some degree, to the satellite-derived PM2.5 used "
    "throughout the main regional panel, and is noted as a limitation of that analysis as "
    "well."
)
print("Fix 6 applied: paragraph 123 tightened.")

# ---------------------------------------------------------------------
# 7. Update the Davao Case Study intro (para 105) so it no longer frames
#    the whole section as just a "standalone validation layer" -- it now
#    has two distinct contributions.
# ---------------------------------------------------------------------
p105 = find_para(text_startswith="The main analysis in this study relies on satellite-derived")
replace_paragraph_text(p105,
    "The main analysis in this study relies on satellite-derived PM2.5 estimates (ACAG "
    "V6.GL.02.04) rather than ground-based monitoring, because ground monitoring in the "
    "Philippines is too sparse and inconsistent to support a 17-region, 10-year panel. Davao "
    "City — one of the few Philippine cities with a multi-year network of government air "
    "quality monitors — offers a chance to check that choice directly: how well does "
    "satellite-derived PM2.5 agree with real, ground-measured PM2.5 at the city level? This "
    "section reports that comparison, and, separately, a broader methodological finding it "
    "surfaced about satellite product-version drift (below) that is relevant beyond this "
    "study. Neither result feeds into, or changes, the main regional panel results."
)
print("Fix 7 applied: Davao Case Study intro (105) reframed to acknowledge two distinct contributions.")

# ---------------------------------------------------------------------
# 8. Light addition to the signpost sentence (para 103, in the
#    Implications-for-policy discussion) so the version-drift finding is
#    visible before the reader reaches the Davao section.
# ---------------------------------------------------------------------
p103 = find_para(text_startswith="Davao City’s regional-average PM2.5 relative to WHO’s Interim Target-3")
replace_paragraph_text(p103,
    "Davao City’s regional-average PM2.5 relative to WHO’s Interim Target-3, how the "
    "satellite-derived estimates used here compare against real ground-station monitoring "
    "from the Environmental Management Bureau Region XI, and a broader satellite "
    "product-version-drift finding this comparison surfaced, are examined directly — "
    "including where the two sources disagree and why — in the Davao City Case Study "
    "subsection immediately below, rather than summarized as a single claim here."
)
print("Fix 8 applied: signpost sentence (103) updated to flag the version-drift finding upfront.")

# ---------------------------------------------------------------------
# 9. Abstract: add a secondary-contribution sentence. This costs nothing
#    against the STS 20-page cap because the abstract is excluded from
#    that count per the official Research Report Guidelines.
# ---------------------------------------------------------------------
abstract = find_para(text_startswith="Background. Fine particulate matter")
old_tail = "and strengthens rather than weakens under a more biologically plausible cumulative-exposure specification."
new_tail = (old_tail + " As a secondary methodological contribution, a ground-station validation "
    "of the underlying satellite PM2.5 product in Davao City identified a substantial "
    "version-drift effect between ACAG product releases (mean absolute error 1.48 vs. 3.65 "
    "µg/m³, a 2.87 µg/m³ average shift) — a caution relevant to any satellite-based ecological "
    "exposure study that spans multiple product versions.")
hit = False
for r in abstract.runs:
    if old_tail in r.text:
        r.text = r.text.replace(old_tail, new_tail)
        hit = True
if not hit:
    raise ValueError("Abstract tail sentence not found for replacement")
print("Fix 9 applied: Abstract Conclusions sentence extended with the version-drift secondary contribution.")

doc.save(PATH)

word_count_after = sum(len(p.text.split()) for p in doc.paragraphs)
print(f"\nWord count before: {word_count_before}")
print(f"Word count after:  {word_count_after}")
print(f"Net word change:   {word_count_after - word_count_before:+d}")
print(f"\nSaved {PATH}")
